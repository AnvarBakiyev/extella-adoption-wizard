$extens("include.py")
include("import docx", ["extella-pip install python-docx"])

def fmt_report_docx(input_path: str = "", records_json: str = "", spec_json: str = "",
                    output_path: str = "", brand_name: str = "", accent: str = "",
                    footer_note: str = "", sections_json: str = "") -> dict:
    """Оформитель отчётов в Word. Тот же контракт, что у fmt_report_pdf — та же спека вида,
    те же входы (записи, файл стадии или готовые разрезы). Разница в назначении:

      PDF  — документ «как есть», его отправляют;
      DOCX — документ, который клиент ДОРАБОТАЕТ: допишет комментарий, добавит абзац,
             отдаст своему заказчику под своим именем. В этом его ценность.

    Кириллица здесь проще, чем в PDF: текст лежит в XML, шрифт подставляет Word —
    искать статический TTF на устройстве не нужно.

    Документ принадлежит КЛИЕНТУ: знаков Extella нет, в шапке имя владельца процесса.
    """
    import json
    import os
    from pathlib import Path
    from datetime import datetime, timezone
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ready = []
    if sections_json and not sections_json.startswith("{{"):
        try:
            ready = [s for s in (json.loads(sections_json) or []) if isinstance(s, dict) and s.get("items")]
        except Exception:
            ready = []

    recs = []
    if records_json and not records_json.startswith("{{"):
        try:
            recs = json.loads(records_json)
        except Exception:
            return {"status": "error", "message": "records_json не разобрался"}
    elif input_path and not input_path.startswith("{{") and Path(input_path).exists():
        try:
            data = json.loads(Path(input_path).read_text(encoding="utf-8"))
        except Exception:
            return {"status": "error", "message": "вход не читается как JSON: " + str(input_path)[:120]}
        recs = data if isinstance(data, list) else (data.get("records") or data.get("rows") or [])
    recs = [r for r in (recs or []) if isinstance(r, dict)]
    if not ready and not recs:
        return {"status": "error", "message": "нет данных для отчёта (ни records_json, ни input_path, ни sections_json)"}

    spec = {}
    if spec_json and not spec_json.startswith("{{"):
        try:
            spec = json.loads(spec_json) or {}
        except Exception:
            spec = {}
    style = spec.get("style") or {}
    accent = accent or style.get("accent") or "#2F6B66"
    brand_name = brand_name or style.get("brand_name") or ""
    footer_note = footer_note or style.get("footer") or ""
    cols = list(recs[0].keys()) if recs else []

    # ── разрезы: готовые или считаем сами (та же логика, что в PDF-оформителе) ──
    sections = list(ready)
    if not sections:
        views = [v for v in (spec.get("views") or []) if isinstance(v, dict) and v.get("group_by") in cols]
        if not views:
            cand = []
            for c in cols:
                vals = [str(r.get(c, "")).strip() for r in recs if str(r.get(c, "")).strip()]
                uniq = len(set(vals))
                if vals and 1 < uniq <= max(2, min(12, len(recs))) and uniq < len(vals):
                    cand.append((uniq, c))
            views = [{"group_by": c, "title": "По полю «" + c + "»"} for _, c in sorted(cand)[:3]]
        for v in views[:4]:
            g = v["group_by"]
            agg = {}
            for r in recs:
                k = str(r.get(g, "")).strip()
                agg[k or "не заполнено"] = agg.get(k or "не заполнено", 0) + 1
            named = {k: n for k, n in agg.items() if k != "не заполнено"}
            items = dict(sorted(named.items(), key=lambda x: -x[1])[:12])
            if agg.get("не заполнено"):
                items["не заполнено"] = agg["не заполнено"]
            sections.append({"title": v.get("title") or g, "items": items})

    hl = spec.get("headline") or {}
    if ready and not recs:
        total = spec.get("total")
        hv = total if total is not None else sum(sum(s["items"].values()) for s in ready)
    else:
        hv = len(recs)
    hlabel = str(hl.get("label") or "записей в отчёте").replace("\n", " ")

    def _rgb(h):
        h = str(h or "").lstrip("#")
        if len(h) != 6:
            h = "2F6B66"
        try:
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        except Exception:
            return RGBColor(0x2F, 0x6B, 0x66)

    A = _rgb(accent)
    SILVER = RGBColor(0x8C, 0x8C, 0x8C)
    INK = RGBColor(0x1A, 0x1A, 0x1A)

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.2)
    base = doc.styles["Normal"]
    base.font.name = "Arial"
    base.font.size = Pt(10.5)
    base.font.color.rgb = INK

    # ── шапка: владелец слева, дата справа ──
    head = doc.add_paragraph()
    r = head.add_run(brand_name or "")
    r.bold = True
    r.font.size = Pt(11)
    tab = head.add_run("\t\t" + datetime.now(timezone.utc).strftime("%d.%m.%Y"))
    tab.font.size = Pt(9)
    tab.font.color.rgb = SILVER

    t = doc.add_paragraph()
    tr = t.add_run(str(spec.get("report") or "Отчёт"))
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = INK
    if spec.get("subtitle"):
        sp = doc.add_paragraph()
        sr = sp.add_run(str(spec["subtitle"]))
        sr.font.size = Pt(10)
        sr.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)

    # ── главное число ──
    hp = doc.add_paragraph()
    hr = hp.add_run(str(hv))
    hr.bold = True
    hr.font.size = Pt(26)
    hr.font.color.rgb = A
    hs = hp.add_run("   " + hlabel)
    hs.font.size = Pt(10.5)

    # ── разрезы таблицами: в Word таблица честнее полос — её можно править и считать ──
    for s in sections:
        hp2 = doc.add_paragraph()
        h2 = hp2.add_run(str(s["title"]).upper())
        h2.bold = True
        h2.font.size = Pt(8.5)
        h2.font.color.rgb = SILVER
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Light List Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Значение"
        hdr[1].text = "Количество"
        for k, v in s["items"].items():
            row = tbl.add_row().cells
            row[0].text = str(k)[:60]
            row[1].text = str(v)
        doc.add_paragraph()

    if footer_note:
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fr = fp.add_run(footer_note)
        fr.font.size = Pt(8)
        fr.font.color.rgb = SILVER

    out = output_path or ("/tmp/report_" + datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S") + ".docx")
    try:
        doc.save(out)
    except Exception as e:
        return {"status": "error", "message": "DOCX не сохранился: " + str(e)[:140]}
    return {"status": "success", "path": out, "bytes": os.path.getsize(out),
            "records": len(recs), "preview": bool(ready and not recs),
            "sections": [s["title"] for s in sections]}
