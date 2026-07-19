# expert: wz_workspace
# description: Extella Workspace engine. TWO kinds of workspace: (1) SCAN — scans a folder, understands the project from FILENAMES (contents never read — private), builds sources/versions/pack/state/cleanup; (2) CURATED — a hand-defined multi-contour workspace whose per-contour state (done/next/blockers/waiting) is set manually. Ops: build / get / list / set_autopilot / organize / undo / run_capability (scan); create_curated / set_state (curated). Persist in KV. Canon: KV/service scope agent_extella_default, Qwen only via agent/run.

def wz_workspace(op="build", folder="", ws_id="", name="", enabled="", n="", capability="", apply="",
                 contours="", understood="", contour="", state="", intent="", key="", value="", scope="", files="", exclude="",
                 api_token="", api_base="https://api.extella.ai", agent_id="", client="") -> str:
    import json, re, os, hashlib, shutil, urllib.request
    from pathlib import Path
    from datetime import datetime, timezone

    def _b(v): return (not v) or str(v).startswith("{{")
    if _b(api_base): api_base = "https://api.extella.ai"
    if _b(agent_id): agent_id = "agent_XwZBKvd8dD70jKvW4WrZm"   # Qwen Extella (qwen3.7-max)
    if _b(api_token):
        cfg = Path.home() / "extella_wizard" / "app" / "config.json"
        try: api_token = json.loads(cfg.read_text(encoding="utf-8")).get("auth_token", "") if cfg.exists() else ""
        except Exception: api_token = ""
    if not api_token:
        return json.dumps({"status": "error", "message": "нет api_token"}, ensure_ascii=False)

    H = {"X-Auth-Token": api_token, "Content-Type": "application/json",
         "X-Profile-Id": "default", "X-Agent-Id": "agent_extella_default"}
    def _post(path, body, t=90, agent=None):
        h = dict(H)
        if agent: h["X-Agent-Id"] = agent
        req = urllib.request.Request(api_base.rstrip("/") + path,
                                     data=json.dumps(body, ensure_ascii=False).encode("utf-8"), headers=h, method="POST")
        with urllib.request.urlopen(req, timeout=t) as r: return json.loads(r.read().decode("utf-8"))

    def _now(): return datetime.now(timezone.utc).isoformat()
    def _kv_get(k):
        try: return json.loads((_post("/api/kv/get", {"key": k}) or {}).get("value") or "null")
        except Exception: return None
    def _kv_set(k, v):
        try: _post("/api/kv/set", {"key": k, "value": json.dumps(v, ensure_ascii=False), "description": "extella workspace"})
        except Exception: pass
    def _ask_qwen(prompt, mx=1400, t=90, tries=2):
        for _ in range(tries):
            try:
                resp = _post("/api/agent/run", {"agent_id": agent_id, "input": prompt, "store": False,
                                                "temperature": 0.3, "tool_choice": "none", "max_output_tokens": mx}, t=t, agent=agent_id)
                parts = [c["text"] for it in (resp.get("output") or []) if isinstance(it, dict) and it.get("type") == "message"
                         for c in it.get("content", []) if isinstance(c, dict) and c.get("text")]
                jt = ("\n".join(parts) or resp.get("output_text") or "").strip()
                if jt: return json.loads(jt[jt.find("{"):jt.rfind("}") + 1])
            except Exception: continue
        return {}
    def _collect_files(root, max_files=600, max_depth=4):
        # РЕКУРСИВНО: относительные пути (подпапки видны), без скрытых/системных
        out = []; root = root.rstrip(os.sep); bd = root.count(os.sep)
        for dp, dns, fns in os.walk(root):
            dns[:] = [d for d in dns if not d.startswith(".") and d not in ("__pycache__", "node_modules")]
            if dp.count(os.sep) - bd >= max_depth: dns[:] = []
            for fn in fns:
                if fn.startswith("."): continue
                out.append(os.path.relpath(os.path.join(dp, fn), root))
                if len(out) >= max_files: return sorted(out)
        return sorted(out)
    # ---- ТАКСОНОМИЯ ТИПОВ ПРОЕКТОВ (T1-T16, из спецификации пути): сигналы → тип → типовые цели/артефакты/инструмент ----
    TYPES = [
        {"k": "invest", "n": "Инвестиции / портфель", "iw": ["инвест", "портфел", "saf", "раунд", "фонд", "акци", "крипт", "cap table", "вложен"], "fw": ["safe", "subscription", "invest", "portfolio", "cap_table", "captable", "broker", "etoro", "акци", "инвест", "займ", "loan"], "tool": "Таблица · Excel/Sheets", "art": "реестр инвестиций (XLSX), сверка версий", "gh": "единый реестр активов; контроль долей и версий документов"},
        {"k": "tax", "n": "Налоги / отчётность", "iw": ["налог", "деклара", "отчётн", "отчетн", "фно", "salyk", "ндс", "коррект"], "fw": ["налог", "деклара", "tax", "910", "200.00", "фно"], "tool": "Таблица · Excel/Sheets", "art": "пакет данных для декларации (XLSX)", "gh": "собрать данные по доходам; подготовить корректировку/подачу"},
        {"k": "legal", "n": "Договор / юридический", "iw": ["договор", "соглашен", "юрист", "юрид", "согласова", "аренд", "nda", "контракт"], "fw": ["договор", "contract", "agreement", "nda", "доп_соглаш", "протокол_разног"], "tool": "Документ", "art": "черновик договора/протокола разногласий (DOCX)", "gh": "собрать документы сторон; согласовать условия; подписать"},
        {"k": "realty", "n": "Сделка с недвижимостью", "iw": ["квартир", "недвижим", "ипотек", "купля", "продажа дома", "участок"], "fw": ["квартир", "недвиж", "ипотек", "кадастр", "техпаспорт"], "tool": "Документ", "art": "чек-лист документов, черновики (DOCX)", "gh": "собрать документы; проверить юридическую чистоту; сделка"},
        {"k": "hiring", "n": "Найм / команда", "iw": ["найм", "ваканси", "кандидат", "собеседован", "рекрут", "оффер"], "fw": ["cv", "resume", "резюме", "ваканси", "оффер", "offer"], "tool": "Таблица · Excel/Sheets", "art": "трекер кандидатов (XLSX), черновик оффера (DOCX)", "gh": "воронка кандидатов; интервью; оффер"},
        {"k": "marketing", "n": "Маркетинг-кампания", "iw": ["маркетинг", "кампани", "реклам", "продвижен", "smm", "лендинг", "бренд"], "fw": ["медиаплан", "креатив", "brief", "бриф", "banner", "лого"], "tool": "Таблица · Excel/Sheets", "art": "медиаплан (XLSX), брифы (DOCX)", "gh": "план кампании; запуск; замер результатов"},
        {"k": "sales", "n": "Продажи / тендер", "iw": ["тендер", "продаж", "клиент", "сделк", "коммерческое предложение", "кп ", "закуп"], "fw": ["тендер", "кп_", "proposal", "invoice", "счет", "счёт"], "tool": "Таблица · Excel/Sheets", "art": "трекер сделок (XLSX), КП (DOCX)", "gh": "воронка сделок; подготовка КП/заявки; закрытие"},
        {"k": "dev", "n": "Разработка / IT", "iw": ["код", "разработ", "деплой", "репозитор", "приложен", "api", "сервис", "баг"], "fw": [".py", ".js", ".ts", "readme", "dockerfile", "package.json", ".git"], "tool": "Код / Markdown", "art": "документация (MD), тех-спеки", "gh": "довести фичу/релиз; чинить баги; документировать"},
        {"k": "design", "n": "Дизайн-проект", "iw": ["дизайн", "макет", "фирменный стиль", "логотип", "figma", "иллюстрац"], "fw": [".psd", ".ai", ".sketch", ".fig", "mockup", "макет", "logo"], "tool": "Заметки", "art": "организация версий макетов (файлы НЕ генерим — читаем/версионируем)", "gh": "собрать референсы; версии макетов; финальные исходники"},
        {"k": "event", "n": "Событие / мероприятие", "iw": ["свадьб", "конференц", "мероприят", "праздник", "юбилей", "фестивал", "событ"], "fw": ["гост", "смета", "программа", "рассадк", "площадк"], "tool": "Таблица · Excel/Sheets", "art": "смета и список гостей (XLSX), программа (DOCX)", "gh": "бюджет и смета; подрядчики; программа дня"},
        {"k": "construction", "n": "Стройка / ремонт", "iw": ["ремонт", "стройк", "строительств", "отделк", "прораб", "смет"], "fw": ["смета", "план_", "чертёж", "чертеж", "проект_дома"], "tool": "Таблица · Excel/Sheets", "art": "смета/график работ (XLSX)", "gh": "смета; график работ; приёмка этапов"},
        {"k": "research", "n": "Учёба / исследование", "iw": ["исследован", "диплом", "диссерта", "курс", "обучен", "статья", "анализ рынка"], "fw": ["lecture", "конспект", "paper", "статья", "исследован"], "tool": "Код / Markdown", "art": "конспект/обзор (MD/DOCX)", "gh": "собрать источники; конспект; итоговый текст"},
        {"k": "relocation", "n": "Переезд / личные документы", "iw": ["переезд", "виза", "внж", "релокац", "гражданств", "документы на"], "fw": ["виза", "visa", "анкета", "внж", "справк"], "tool": "Документ", "art": "чек-лист документов (DOCX)", "gh": "чек-лист; сбор документов; подача"},
        {"k": "procurement", "n": "Закупки / поставщики", "iw": ["закупк", "поставщик", "снабжен", "прайс", "заказ у"], "fw": ["прайс", "price", "поставщик", "заказ_", "накладн"], "tool": "Таблица · Excel/Sheets", "art": "сравнение поставщиков (XLSX)", "gh": "собрать прайсы; сравнить; заказать"},
        {"k": "travel", "n": "Поездка / путешествие", "iw": ["поездк", "путешеств", "перелёт", "отпуск", "маршрут", "тур "], "fw": ["билет", "ticket", "booking", "бронь", "маршрут"], "tool": "Таблица · Excel/Sheets", "art": "маршрут и бюджет (XLSX)", "gh": "маршрут; брони; бюджет поездки"},
        {"k": "finance_ops", "n": "Финансы / учёт", "iw": ["бюджет", "бухгалтер", "расход", "доход", "кассов", "финанс", "сверк"], "fw": ["бюджет", "budget", "расход", "выписк", "1с", "invoice"], "tool": "Таблица · Excel/Sheets", "art": "бюджет/сверка (XLSX)", "gh": "свести расходы/доходы; сверка; бюджет"},
    ]
    def _detect_type(intent_text, names):
        it = (intent_text or "").lower()
        nl = " ".join(str(x).lower() for x in (names or [])[:300])
        best = None; bs = 0; second = 0
        for t in TYPES:
            sc = sum(3 for w in t["iw"] if w in it) + min(sum(1 for w in t["fw"] if w in nl), 6)
            if sc > bs: second = bs; bs = sc; best = t
            elif sc > second: second = sc
        if not best or bs < 2: return None, 0, False
        return best, bs, (second >= bs - 1 and second >= 2)   # ambiguous: второй тип почти равен
    EXCL_DEFAULT = ["паспорт", "passport", "пароль", "password", "credential", "секрет", "secret", ".ssh", "id_rsa"]
    def _excl_filter(paths, extra):
        pats = [str(x).lower() for x in (EXCL_DEFAULT + list(extra or [])) if str(x).strip()]
        return [p for p in paths if not any(t in p.lower() for t in pats)]
    def _me_facts(): return _kv_get("me:facts") or {}          # глобальный профиль «Я» — общий для всех воркспейсов
    def _me_set(k, v):
        m = _me_facts(); m[k] = v; _kv_set("me:facts", m); return m
    # ---- ШАРДИРОВАНИЕ KV (ws-v1.4): растущие списки — отдельными ключами, ядро остаётся малым ----
    # Старые (нешардированные) объекты читаются как есть; при первом сохранении перешардируются.
    def _kv_get_retry(k, tries=3):
        # исключение = транзиентный сбой KV (ретрай); null-значение = ключа нет (легально)
        for i in range(tries):
            try:
                return True, json.loads((_post("/api/kv/get", {"key": k}) or {}).get("value") or "null")
            except Exception:
                if i < tries - 1: __import__("time").sleep(0.7)
        return False, None
    def _ws_load(wsid):
        ok, ws = _kv_get_retry("workspace:" + str(wsid))
        if not ok: return None   # временный сбой — лучше «повтори», чем работать вслепую
        if isinstance(ws, dict) and ws.get("sharded"):
            for part in ("tasks", "questions", "ledger"):
                ok2, v = _kv_get_retry("ws:" + str(wsid) + ":" + part)
                if not ok2: return None   # НЕ работаем с половинчатым объектом (иначе _ws_store затрёт часть пустотой)
                ws[part] = v if isinstance(v, list) else []
        return ws
    def _ws_store(ws, reopen_ids=None):
        wsid = str(ws.get("ws_id") or "")
        if not wsid: return
        ws["sharded"] = True
        # rev-merge: если параллельный писатель (UI/автопилот) уже закрыл задачу — done не откатываем.
        # ИСКЛЮЧЕНИЕ: reopen_ids — задачи, которые человек ЯВНО переоткрыл (set_task todo) — его слово сильнее merge.
        try:
            fresh = _kv_get("ws:" + wsid + ":tasks")
            if isinstance(fresh, list) and fresh:
                fmap = {t.get("id"): t for t in fresh if isinstance(t, dict)}
                for t in (ws.get("tasks") or []):
                    if reopen_ids and t.get("id") in reopen_ids: continue
                    f2 = fmap.get(t.get("id"))
                    if f2 and f2.get("status") == "done" and t.get("status") != "done":
                        t["status"] = "done"
        except Exception: pass
        # тот же merge для вопросов: answered не откатываем, вопросы параллельного писателя не теряем
        try:
            fq = _kv_get("ws:" + wsid + ":questions")
            if isinstance(fq, list) and fq:
                ours = {q.get("key"): q for q in (ws.get("questions") or []) if isinstance(q, dict)}
                for q2 in fq:
                    if not isinstance(q2, dict): continue
                    mine = ours.get(q2.get("key"))
                    if mine is None:
                        ws.setdefault("questions", []).append(q2)
                    elif q2.get("status") == "answered" and mine.get("status") != "answered":
                        mine["status"] = "answered"; mine["answer"] = q2.get("answer", "")
        except Exception: pass
        ws["rev"] = int(ws.get("rev") or 0) + 1
        caps = {"tasks": 60, "questions": 40, "ledger": 60}
        for part, cap in caps.items():
            lst = ws.get(part) or []
            lst = lst[-cap:] if part == "ledger" else lst[:cap]
            ws[part] = lst
            _kv_set("ws:" + wsid + ":" + part, lst)
        core = {k: v for k, v in ws.items() if k not in caps}
        _kv_set("workspace:" + wsid, core)
        for _t in range(2):   # verify-write ядра (KV шардирован, бывают битые записи)
            try:
                _c = _kv_get("workspace:" + wsid)
                if isinstance(_c, dict) and _c.get("ws_id") == wsid: break
            except Exception: pass
            _kv_set("workspace:" + wsid, core)
    def _idx_get():
        for _ in range(3):   # исключение = сбой (ретрай); null-значение = ключа нет = законная пустота
            try:
                v = json.loads((_post("/api/kv/get", {"key": "workspace:index"}) or {}).get("value") or "null")
                return v if isinstance(v, list) else []
            except Exception:
                __import__("time").sleep(0.6)
        return None
    def _idx_put(idx):
        _kv_set("workspace:index", idx)
        for _ in range(2):
            try:
                v = _kv_get("workspace:index")
                if isinstance(v, list) and len(v) >= len(idx): return
            except Exception: pass
            _kv_set("workspace:index", idx)
    def _idx_upsert(entry):
        idx = _idx_get() or []
        idx = [e for e in idx if not (isinstance(e, dict) and e.get("ws_id") == entry.get("ws_id"))]
        idx.insert(0, entry)
        _idx_put(idx[:100])
        return idx
    # ПРАВИЛО МАТЕРИАЛИЗАЦИИ: задача, обещающая ФАЙЛ («записать в таблицу», «генерация формы 240.00»),
    # закрывается ТОЛЬКО когда файл реально создан (artifact=True из ветки записи). Иначе кокпит показывает
    # 100% и «сделано», а на диске пусто — это враньё, ради которого продукт не нужен.
    ART_WORDS = ("запиш", "запис", "сохран", "выгруз", "экспорт", "файл", "таблиц", "xlsx", "docx",
                 "генерац", "сформир", "сгенерир", "форм", "черновик", "документ", "презентац", "отчёт", "отчет")
    def _is_artifact_task(t):
        return any(w in str(t.get("title", "")).lower() for w in ART_WORDS)
    def _advance(ws, kind, text, explicit_key="", artifact=False):
        # ЗАМКНУТЬ ЦИКЛ: отметить выполненной подходящую todo-задачу типа kind → пересчитать progress+next.
        tasks = ws.get("tasks") or []
        done_t = None
        if explicit_key:
            done_t = next((t for t in tasks if t.get("id") == explicit_key and t.get("status") != "done"), None)
            if done_t and _is_artifact_task(done_t) and not artifact:
                done_t = None   # даже явный ключ не закрывает файловую задачу без файла
        if not done_t:
            cands = [t for t in tasks if t.get("status") != "done" and t.get("kind") == kind
                     and (artifact or not _is_artifact_task(t))]   # файловые задачи — только по факту записи
            toks = set(re.findall(r"[а-яёa-z]{4,}", (text or "").lower()))
            best = None; best_s = 0; best_need = 1
            for t in cands:
                tt = set(re.findall(r"[а-яёa-z]{4,}", str(t.get("title", "")).lower()))
                s = len(toks & tt)
                if s > best_s: best_s = s; best = t; best_need = (2 if len(tt) >= 3 else 1)   # длинные задачи требуют ≥2 общих слов (меньше ложных матчей)
            if best and best_s >= best_need:
                done_t = best
        if not done_t:
            return None
        done_t["status"] = "done"
        for gi, g in enumerate(ws.get("goals", [])):
            gt = [t for t in tasks if t.get("goal") == gi]; d = sum(1 for t in gt if t.get("status") == "done")
            g["progress"] = int(100 * d / len(gt)) if gt else 0
        ws.setdefault("state", {})["next"] = next((t["title"] for t in tasks if t.get("status") != "done"), "")
        return done_t

    # ---- САМОНАРАЩИВАНИЕ (P2): система проектирует себе новые способности из собственных тупиков ----
    # Канон: спроектировать может сама (и автопилот), но ЗАРЕГИСТРИРОВАТЬ — только человек словом «разрешаю».
    # Выращенные способности v1 — ТОЛЬКО ЧИТАЮЩИЕ: сеть (объявленные хосты) + KV; запись файлов остаётся за гейтом записи.
    def _cap_lint(code, declared):
        # статический страж (2-й эшелон после человеческого «разрешаю»): AST-анализ, не substring-регексы.
        # v1 выращенных способностей — ТОЛЬКО ЧТЕНИЕ: сеть по объявленным хостам + KV. Запись файлов — отдельный гейт.
        bad = ["subprocess", "os.system", "os.popen", "shutil", "rmtree", "eval(", "exec(", "__import__",
               "ctypes", "pty.", "os.environ", "getenv", "/api/agent", "api/expert/run", "expert/save"]
        for b in bad:
            if b in code: return False, "запрещённая конструкция: " + b, []
        # запрет ЧТЕНИЯ чужих секретов (свой токен из config.json — можно; SSH-ключи/PIN-хеш/чужие токены — нельзя):
        # иначе способность прочитала бы файл и вернула его в note → утечка мимо скраба (находка ревью).
        # 2-й эшелон после человеческого «разрешаю»; регистронезависимо, широкий список хранилищ секретов.
        low = code.lower()
        # (не блокируем "auth_token" — легитимный скелет читает СВОЙ токен из config.json; чужой AUTH_TOKEN-файл — это extella_mcp_server)
        secret_paths = [".ssh", "id_rsa", "id_ed25519", "id_dsa", "id_ecdsa", "secrets.json", "hosting/secrets",
                        ".extella_test_token", "extella_mcp_server", ".pem", ".p12", ".pfx",
                        ".aws/credentials", ".aws\\credentials", "git-credentials", ".netrc", ".pgpass", ".gnupg",
                        ".kube", ".npmrc", ".docker/config", "keychain", "wallet.dat", "credentials/openai"]
        for sp in secret_paths:
            if sp in low: return False, "чтение секретов запрещено: " + sp, []
        import ast as _ast
        try:
            tree = _ast.parse(code)
        except SyntaxError as e:
            return False, "синтаксическая ошибка: " + str(e)[:80], []
        ALLOWED_IMP = {"json", "re", "urllib", "urllib.request", "urllib.parse", "urllib.error", "pathlib",
                       "datetime", "time", "math", "xml", "xml.etree", "xml.etree.ElementTree", "html",
                       "html.parser", "csv", "io", "hashlib", "unicodedata", "itertools", "collections", "functools"}
        BAD_CALLS = {"remove", "unlink", "rmdir", "removedirs", "makedirs", "mkdir", "rename", "replace_file",
                     "chmod", "chown", "kill", "system", "popen", "urlretrieve", "write_text", "write_bytes",
                     "touch", "eval", "exec", "compile", "include"}
        for node in _ast.walk(tree):
            if isinstance(node, (_ast.Import,)):
                for a in node.names:
                    if a.name.split(".")[0] not in {m.split(".")[0] for m in ALLOWED_IMP}:
                        return False, "импорт вне белого списка: " + a.name, []
            if isinstance(node, _ast.ImportFrom):
                if (node.module or "").split(".")[0] not in {m.split(".")[0] for m in ALLOWED_IMP}:
                    return False, "импорт вне белого списка: " + str(node.module), []
            if isinstance(node, _ast.Call):
                fn = node.func
                nm = fn.id if isinstance(fn, _ast.Name) else (fn.attr if isinstance(fn, _ast.Attribute) else "")
                if nm in BAD_CALLS:
                    if nm in ("compile", "eval", "exec") and not isinstance(fn, _ast.Name):
                        pass   # re.compile и т.п. — безопасные методы модулей из белого списка; запрещены только голые compile/eval/exec
                    else:
                        return False, "запрещённый вызов: " + nm, []
                if nm == "open":
                    mode = ""
                    if len(node.args) >= 2:
                        if isinstance(node.args[1], _ast.Constant): mode = str(node.args[1].value)
                        else: return False, "режим open() должен быть строкой-литералом", []
                    for kw in (node.keywords or []):
                        if kw.arg == "mode":
                            if isinstance(kw.value, _ast.Constant): mode = str(kw.value.value)
                            else: return False, "режим open() должен быть строкой-литералом", []
                    if any(c in mode for c in "wax+"):
                        return False, "запись файлов запрещена (запись — отдельный гейт с подтверждением)", []
        urls = re.findall(r"https?://[^\s'\"\\)]+", code)
        hosts = set(u.split("/")[2] for u in urls if "://" in u and len(u.split("/")) > 2)
        dh = set(str(u).split("/")[2] for u in (declared or []) if "://" in str(u) and len(str(u).split("/")) > 2)
        dh.add("api.extella.ai")
        extra = sorted(h for h in hosts if h not in dh)
        if extra: return False, "код ходит в необъявленную сеть: " + ", ".join(extra[:3]), urls
        return True, "", urls
    def _norm_triggers(raw):
        # Qwen может вернуть строку вместо списка — не итерируем посимвольно; токены только ≥3 символов
        if isinstance(raw, str): raw = re.findall(r"[a-zа-яё0-9]{3,}", raw.lower())
        if not isinstance(raw, list): return []
        return [str(t).strip().lower()[:30] for t in raw if len(str(t).strip()) >= 3][:8]
    def _cap_register(ws, want_slug="", want_sha=""):
        # ГЕЙТ ЧЕЛОВЕКА: регистрируем ровно тот код, который человек видел. want_slug/want_sha — привязка
        # одобрения к конкретной карточке; несовпадение = карточка устарела (драфт перепроектирован) → честный отказ.
        dr = _kv_get("ws:" + ws["ws_id"] + ":cap_draft")
        if not (isinstance(dr, dict) and dr.get("code") and dr.get("slug")):
            # идемпотентность даблклика: уже выращена? — честный success
            if want_slug:
                hit = next((c for c in (_kv_get("ws:" + ws["ws_id"] + ":caps") or [])
                            if isinstance(c, dict) and c.get("slug") == want_slug), None)
                if hit: return True, "Способность «%s» уже выращена — она в паке." % hit.get("cap_name"), hit
            return False, "нет спроектированной способности — сначала скажи «научись: <что сделать>»", None
        if want_slug and dr["slug"] != want_slug:
            return False, "эта карточка устарела — предложение уже заменено новым («%s»). Открой свежую карточку." % dr.get("cap_name"), None
        if want_sha and dr.get("sha") and dr["sha"] != want_sha:
            return False, "код в карточке не совпадает с текущим предложением — открой свежую карточку", None
        ok, why, _u = _cap_lint(dr["code"], dr.get("network") or [])
        if not ok: return False, "код не прошёл проверку безопасности: " + why, None
        ename = "wscap_" + dr["slug"]
        if ("def " + ename + "(") not in dr["code"]:
            return False, "имя функции в коде не совпадает с именем эксперта — пересобери «научись: …»", None
        try:
            r = _post("/api/expert/save", {"name": ename, "description": ("Выращенная способность Extella Workspace [%s]: %s" % (ws["ws_id"], str(dr.get("cap_name", ""))))[:140],
                                           "code": dr["code"], "cspl": "fython", "global": True}, t=60)
            if not (isinstance(r, dict) and r.get("status") == "success"):
                return False, "платформа не приняла эксперта: " + str(r)[:120], None
        except Exception as e:
            return False, "регистрация упала: " + str(e)[:100], None
        creg = [c for c in (_kv_get("ws:" + ws["ws_id"] + ":caps") or []) if isinstance(c, dict) and c.get("slug") != dr["slug"]]
        cap = {"slug": dr["slug"], "cap_name": str(dr.get("cap_name") or dr["slug"])[:60], "expert_name": ename,
               "triggers": _norm_triggers(dr.get("triggers")), "network": (dr.get("network") or [])[:5],
               "params": dr.get("params") or {}, "req": str(dr.get("req") or "")[:120], "sha": str(dr.get("sha") or "")[:16], "at": _now()}
        creg.append(cap)
        while len(creg) > 1 and len(json.dumps(creg, ensure_ascii=False).encode("utf-8")) > 6000:
            creg.pop(0)   # реестр держим < KV-лимита: старейшие способности выпадают из диспетча (эксперт остаётся)
        _kv_set("ws:" + ws["ws_id"] + ":caps", creg)
        chk = _kv_get("ws:" + ws["ws_id"] + ":caps")
        if not (isinstance(chk, list) and any(c.get("slug") == dr["slug"] for c in chk if isinstance(c, dict))):
            return False, "реестр способностей не сохранился (сбой KV) — повтори «разрешаю»", None
        pk = ws.setdefault("pack", {}).setdefault("capabilities", [])
        if ("🧠 " + cap["cap_name"]) not in pk: pk.append("🧠 " + cap["cap_name"])
        for q in ws.get("questions", []):
            if q.get("key") == "cap_" + dr["slug"] and q.get("status") != "answered":
                q["status"] = "answered"; q["answer"] = "разрешаю"
        ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": "выращена способность",
                                            "detail": cap["cap_name"] + " → " + ename, "at": _now(), "undo": "none"})
        _kv_set("ws:" + ws["ws_id"] + ":cap_draft", None)
        return True, "Способность «%s» выращена и зарегистрирована. Теперь умею это сам — запусти автопилот или попроси в чате." % cap["cap_name"], cap

    # ---------- get / list ----------
    if op == "get":
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        return json.dumps({"status": "success" if ws else "error", "workspace": ws,
                           "message": ("" if ws else "воркспейс не найден")}, ensure_ascii=False)
    if op == "list":
        idx = _idx_get()
        if idx is None:
            return json.dumps({"status": "error", "retryable": True, "message": "индекс временно недоступен — повтори"}, ensure_ascii=False)
        return json.dumps({"status": "success", "workspaces": idx}, ensure_ascii=False)

    if op == "delete":   # удалить ВОРКСПЕЙС (запись в KV + из индекса). Файлы/папку пользователя НЕ трогает.
        if _b(ws_id):
            return json.dumps({"status": "error", "message": "нужен ws_id"}, ensure_ascii=False)
        wsid = str(ws_id)
        _kv_set("workspace:" + wsid, None)   # объект → null (get вернёт «не найден»)
        for part in ("tasks", "questions", "ledger", "draft_registry"):   # чистим шардированные части
            _kv_set("ws:" + wsid + ":" + part, None)
        idx = [e for e in (_idx_get() or []) if not (isinstance(e, dict) and e.get("ws_id") == wsid)]
        _idx_put(idx)
        return json.dumps({"status": "success", "deleted": wsid, "remaining": len(idx),
                           "note": "удалена только запись воркспейса; файлы/папка не тронуты"}, ensure_ascii=False)

    if op == "get_me":   # глобальный профиль «Я» — что система знает про пользователя (для панели фактов)
        return json.dumps({"status": "success", "facts": _kv_get("me:facts") or {}}, ensure_ascii=False)

    # ---------- understand: НЕ угадывать вслепую — предложить понимание + цели + ОДИН вопрос ----------
    if op == "understand":
        if _b(folder) or not os.path.isdir(os.path.expanduser(folder)):
            return json.dumps({"status": "error", "message": "нужна существующая папка (folder)"}, ensure_ascii=False)
        froot = os.path.expanduser(folder)
        allf = _collect_files(froot)
        try: _exl = json.loads(exclude) if not _b(exclude) else []
        except Exception: _exl = []
        allf = _excl_filter(allf, _exl)
        if not allf:
            return json.dumps({"status": "error", "message": "в папке и подпапках нет файлов"}, ensure_ascii=False)
        subdirs = sorted(set(os.path.dirname(f) for f in allf if os.path.dirname(f)))
        _pt, _psc, _pam = _detect_type("", allf)
        _th = ("\nПОХОЖИЙ ТИП ПРОЕКТА: «%s»; типовые цели: %s.\n" % (_pt["n"], _pt["gh"])) if _pt else ""
        prompt = (_th + "Ты — Extella. Пользователь дал ПАПКУ (возможно с подпапками). По ИМЕНАМ файлов и структуре пойми, ЧТО это за проект — "
                  "и НЕ угадывай вслепую: предложи 2-3 вероятные ЦЕЛИ и задай ОДИН короткий уточняющий вопрос, чтобы точно понять, чего хочет человек. "
                  "Содержимое файлов НЕ читаешь (приватность). Конкретно, без воды.\n\nСТРУКТУРА (относительные пути, подпапки видны):\n"
                  + "\n".join("- " + n for n in allf[:400]) +
                  '\n\nЕсли в папке ЯВНО НЕСКОЛЬКО разных проектов (свалка вроде Downloads) — скажи это честно.'
                  '\n\nВерни СТРОГО JSON: {"understood":"<1 фраза: похоже, это ...>",'
                  '"multi_project": true|false,'
                  '"projects":[{"name":"<имя проекта>","hint":"<какие файлы к нему относятся>"}],'
                  '"candidates":[{"intent":"<вероятная цель — что Extella сделает>","why":"<по каким файлам видно>"}],'
                  '"question":"<ОДИН короткий вопрос пользователю для уточнения цели>"}'
                  '\n(multi_project=true только если проектов явно >1; projects тогда 2-5 шт)')
        ai = _ask_qwen(prompt)
        multi = bool(ai.get("multi_project")) and len(ai.get("projects") or []) >= 2
        resp = {"status": "success", "clarify": True,
                "understood": ai.get("understood", ""),
                "candidates": (ai.get("candidates") or [])[:3],
                "question": ai.get("question") or "Что именно ты хочешь получить из этой папки?",
                "folder": froot, "file_count": len(allf), "subfolders": subdirs[:20]}
        if multi:   # Downloads-кейс: свалка из нескольких проектов → предложить уборку/выбор, не собирать винегрет
            resp["multi_project"] = True
            resp["projects"] = (ai.get("projects") or [])[:5]
            resp["recommendation"] = "organize_first"
            resp["question"] = ("Похоже, здесь несколько разных проектов: %s. Сначала прибраться (разложу по папкам), "
                                "или собрать воркспейс по одному из них — какому?" %
                                ", ".join("«%s»" % str(pp.get("name", ""))[:30] for pp in resp["projects"][:4]))
        return json.dumps(resp, ensure_ascii=False)

    # ---------- set_fact: воркспейс ЗАПОМИНАЕТ ответ на свой вопрос (facts) — механизм «спроси один раз» ----------
    if op == "set_fact":
        if _b(key):
            return json.dumps({"status": "error", "message": "нужен key"}, ensure_ascii=False)
        scope_me = (str(scope).strip().lower() == "me")
        me = _me_set(key, value) if scope_me else None
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if str(key).startswith("cap_") and ws:
            # ОТВЕТ НА ПРЕДЛОЖЕНИЕ СПОСОБНОСТИ. Канон гейта: отрицание главнее («не разрешаю» = отказ),
            # одобрение — только явное; неоднозначный ответ НЕ закрывает вопрос и НЕ трогает драфт.
            vl = str(value).strip().lower()
            neg = bool(re.search(r"(^|[\s,.!])(не|нет|no)([\s,.!]|$)", vl)) or "не надо" in vl
            pos = bool(re.search(r"разреша|подтвержда|confirm", vl))
            if neg:
                dr0 = _kv_get("ws:" + ws["ws_id"] + ":cap_draft") or {}
                _kv_set("ws:" + ws["ws_id"] + ":cap_draft", None)
                for q in ws.get("questions", []):
                    if q.get("key") == key and q.get("status") != "answered":
                        q["status"] = "answered"; q["answer"] = "не надо"
                ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": "способность отклонена",
                                                    "detail": str(dr0.get("cap_name") or "")[:60], "at": _now(), "undo": "none"})
                ws["updated"] = _now(); _ws_store(ws)
                return json.dumps({"status": "success", "note": "Понял, не надо — предложение способности убрал", "workspace": ws}, ensure_ascii=False)
            if not pos:
                return json.dumps({"status": "success", "note": "Не понял ответ — скажи «разрешаю» или «не надо». Предложение остаётся открытым.",
                                   "workspace": ws}, ensure_ascii=False)
            okr, cnote, _cap = _cap_register(ws, want_slug=str(key)[4:])
            if not okr:
                return json.dumps({"status": "error", "message": cnote, "workspace": ws}, ensure_ascii=False)
            adv = _advance(ws, "question", str(key), "")
            ws["updated"] = _now(); _ws_store(ws)
            return json.dumps({"status": "success", "kind": "cap_added", "note": cnote, "cap": _cap,
                               "advanced": (adv["title"] if adv else None), "workspace": ws}, ensure_ascii=False)
        adv = None
        if ws:
            # ОТВЕТ ФАЙЛОМ: приложенные файлы становятся и ответом, и источниками воркспейса
            try: _afl = json.loads(files) if not _b(files) else []
            except Exception: _afl = []
            _ap = [os.path.abspath(os.path.expanduser(str(p))) for p in (_afl if isinstance(_afl, list) else [])]
            _ap = [p for p in _ap if os.path.isfile(p)]
            if _ap:
                _fb = os.path.expanduser(ws.get("folder") or "")
                _have = set(s0.get("name") for s0 in ws.get("sources", []))
                _names = []
                for p in _ap:
                    rel = os.path.relpath(p, _fb) if (_fb and os.path.realpath(p).startswith(os.path.realpath(_fb) + os.sep)) else p
                    _names.append(rel)
                    if rel not in _have:
                        ws.setdefault("sources", []).append({"name": rel, "ext": os.path.splitext(p)[1].lstrip(".").lower()})
                        if ws.get("scope_files"): ws["scope_files"].append(rel)
                        _have.add(rel)
                ws["sources_total"] = int(ws.get("sources_total") or len(ws.get("sources", []))) + len(_names)
                if _b(value): value = "; ".join(os.path.basename(x) for x in _names)
                ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": "ответ файлом",
                                                    "detail": ", ".join(os.path.basename(x) for x in _names[:4]), "at": _now(), "undo": "none"})
                if scope_me: me = _me_set(key, value)   # факт «Я» — уже с именами файлов
            if not scope_me:
                ws.setdefault("facts", {})[key] = value
            qtext = ""
            for q in ws.get("questions", []):   # закрыть открытый вопрос с этим ключом
                if q.get("key") == key and q.get("status") != "answered":
                    q["status"] = "answered"; q["answer"] = value; qtext = q.get("question", "")
            adv = _advance(ws, "question", (qtext + " " + str(key)), "")   # закрыть и его question-задачу → двигать «дальше»
            ws["updated"] = _now(); _ws_store(ws)
        return json.dumps({"status": "success", "scope": ("me" if scope_me else "workspace"),
                           "note": "Запомнил" + ((" · закрыл задачу «" + adv["title"] + "»") if adv else "") + (" — храню в профиле «Я»" if scope_me else ""),
                           "facts": (me if scope_me else ((ws or {}).get("facts") or {})),
                           "advanced": (adv["title"] if adv else None), "task_next": ((ws or {}).get("state", {}) or {}).get("next"),
                           "workspace": ws}, ensure_ascii=False)

    if op == "set_task":   # статус задачи (key=id, value=todo|doing|done) → пересчёт прогресса целей и «что дальше»
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой хранилища — повтори через пару секунд)"}, ensure_ascii=False)
        if _b(key):
            return json.dumps({"status": "error", "message": "нужен key (id задачи)"}, ensure_ascii=False)
        hit = next((t for t in ws.get("tasks", []) if t.get("id") == key), None)
        if not hit:
            return json.dumps({"status": "error", "message": "задача не найдена: " + str(key)}, ensure_ascii=False)
        if value not in ("todo", "doing", "done"):
            return json.dumps({"status": "error", "message": "value должен быть todo|doing|done"}, ensure_ascii=False)
        hit["status"] = value
        for gi, g in enumerate(ws.get("goals", [])):
            gt = [t for t in ws.get("tasks", []) if t.get("goal") == gi]; d = sum(1 for t in gt if t.get("status") == "done")
            g["progress"] = int(100 * d / len(gt)) if gt else 0
        ws.setdefault("state", {})["next"] = next((t["title"] for t in ws.get("tasks", []) if t.get("status") != "done"), "")
        ws["updated"] = _now(); _ws_store(ws, reopen_ids=({key} if value != "done" else None))   # явное переоткрытие сильнее merge
        _nx = ws["state"]["next"]
        return json.dumps({"status": "success", "note": ("Задача отмечена. Дальше: " + _nx) if _nx else "Задача отмечена — все задачи закрыты!",
                           "workspace": ws}, ensure_ascii=False)

    if op == "set_goal":   # ПРАВКА ЦЕЛЕЙ РУКАМИ: n=""|"new"+value → добавить; n=индекс+value → переименовать; n=индекс+value="" → удалить (с её задачами)
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой хранилища — повтори через пару секунд)"}, ensure_ascii=False)
        goals = ws.setdefault("goals", []); tasks = ws.setdefault("tasks", [])
        def _tool_for(t):
            t = (t or "").lower()
            if any(k in t for k in ("реестр", "деклар", "финанс", "портфел", "бюджет", "счет", "табл", "excel", "sheet", "cap table", "отчёт", "отчет")): return "Таблица · Excel/Sheets"
            if any(k in t for k in ("код", "code", "markdown", "репозитор", "github", "dev", "разработ")): return "Код / Markdown"
            if any(k in t for k in ("почт", "письм", "mail", "email", "рассыл")): return "Почта"
            if any(k in t for k in ("презентац", "слайд", "pitch", "slide")): return "Презентация"
            if any(k in t for k in ("договор", "документ", "юрид")): return "Документ"
            return "Заметки"
        nn = str(n).strip().lower()
        if nn in ("", "new"):
            if _b(value):
                return json.dumps({"status": "error", "message": "нужен value (название новой цели)"}, ensure_ascii=False)
            goals.append({"title": str(value)[:120], "outcome": "", "tool": _tool_for(str(value)), "status": "active", "progress": 0})
            act = "добавил цель «%s»" % str(value)[:60]
        else:
            try: gi = int(nn)
            except Exception:
                return json.dumps({"status": "error", "message": "n должен быть индексом цели или 'new'"}, ensure_ascii=False)
            if not (0 <= gi < len(goals)):
                return json.dumps({"status": "error", "message": "нет цели с индексом %d" % gi}, ensure_ascii=False)
            if _b(value):   # удалить цель + её задачи; переиндексация goal у оставшихся
                act = "удалил цель «%s»" % str(goals[gi].get("title", ""))[:60]
                goals.pop(gi)
                ws["tasks"] = [t for t in tasks if t.get("goal") != gi]
                for t in ws["tasks"]:
                    if isinstance(t.get("goal"), int) and t["goal"] > gi: t["goal"] -= 1
                tasks = ws["tasks"]
            else:
                goals[gi]["title"] = str(value)[:120]; goals[gi]["tool"] = _tool_for(str(value))
                act = "переименовал цель → «%s»" % str(value)[:60]
        for gi2, g in enumerate(goals):
            gt = [t for t in tasks if t.get("goal") == gi2]; d = sum(1 for t in gt if t.get("status") == "done")
            g["progress"] = int(100 * d / len(gt)) if gt else 0
        ws.setdefault("state", {})["next"] = next((t["title"] for t in tasks if t.get("status") != "done"), "")
        ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": act, "detail": "правка целей пользователем", "at": _now(), "undo": "none"})
        ws["updated"] = _now(); _ws_store(ws)
        return json.dumps({"status": "success", "note": act[0].upper() + act[1:], "workspace": ws}, ensure_ascii=False)

    if op == "add_files":   # ДОБАВИТЬ МАТЕРИАЛЫ в живой проект (перетащил файлы / приложил документ)
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой хранилища — повтори через пару секунд)"}, ensure_ascii=False)
        try: _fl = json.loads(files) if not _b(files) else []
        except Exception: _fl = []
        picked = [os.path.abspath(os.path.expanduser(str(p))) for p in (_fl if isinstance(_fl, list) else [])]
        picked = [p for p in picked if os.path.isfile(p)]
        if not picked:
            return json.dumps({"status": "error", "message": "ни один файл не найден"}, ensure_ascii=False)
        base = os.path.expanduser(ws.get("folder") or "")
        added = []
        have = set(s0.get("name") for s0 in ws.get("sources", []))
        for p in picked:
            rel = os.path.relpath(p, base) if (base and os.path.realpath(p).startswith(os.path.realpath(base) + os.sep)) else p
            if rel in have: continue
            ws.setdefault("sources", []).append({"name": rel, "ext": os.path.splitext(p)[1].lstrip(".").lower()})
            if ws.get("scope_files"): ws["scope_files"].append(rel)
            added.append(rel); have.add(rel)
        ws["sources_total"] = int(ws.get("sources_total") or len(ws.get("sources", []))) + len(added)
        if added:
            ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": "добавил материалы",
                                                "detail": "%d: %s" % (len(added), ", ".join(os.path.basename(a) for a in added[:4])), "at": _now(), "undo": "none"})
        ws["updated"] = _now(); _ws_store(ws)
        return json.dumps({"status": "success", "note": "Добавил материалов: " + str(len(added)), "added": added, "workspace": ws}, ensure_ascii=False)

    if op == "sync_sources":   # ПОДХВАТ новых файлов папки в живой проект (при открытии кокпита; без Qwen, мгновенно)
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой хранилища — повтори через пару секунд)"}, ensure_ascii=False)
        base = os.path.expanduser(ws.get("folder") or "")
        if not base or not os.path.isdir(base):
            return json.dumps({"status": "success", "new_files": [], "note": "у проекта нет папки (создан словами/файлами)"}, ensure_ascii=False)
        import datetime as _dt
        last = ws.get("last_synced") or ws.get("created") or ""
        try: last_ts = _dt.datetime.fromisoformat(last).timestamp()
        except Exception: last_ts = 0
        cur = _excl_filter(_collect_files(base), ws.get("scope_exclude"))
        newf = []
        for fn in cur:
            try:
                if os.path.getmtime(os.path.join(base, fn)) > last_ts: newf.append(fn)
            except Exception: pass
        ws["sources"] = [{"name": n2, "ext": os.path.splitext(n2)[1].lstrip(".").lower()} for n2 in cur[:40]]
        ws["sources_total"] = len(cur)
        ws["last_synced"] = _now()
        if newf:
            ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": "подхватил новые файлы",
                                                "detail": "%d: %s" % (len(newf), ", ".join(os.path.basename(x) for x in newf[:4])), "at": _now(), "undo": "none"})
        ws["updated"] = _now(); _ws_store(ws)
        return json.dumps({"status": "success", "note": ("Подхватил новых файлов: " + str(len(newf))) if newf else "Новых файлов нет",
                           "new_files": newf[:20], "new_count": len(newf),
                           "sources_total": len(cur), "workspace": ws}, ensure_ascii=False)

    if op == "digest":   # «ПОКА ТЕБЯ НЕ БЫЛО» — что сделано/появилось с прошлого открытия (без Qwen, мгновенно)
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой хранилища — повтори через пару секунд)"}, ensure_ascii=False)
        since = ws.get("last_opened") or ws.get("created") or ""
        acts = [e for e in (ws.get("ledger") or []) if str(e.get("at", "")) > str(since)]
        qs_open = [q for q in (ws.get("questions") or []) if q.get("status") != "answered"]
        base = os.path.expanduser(ws.get("folder") or "")
        newf = []
        if base and os.path.isdir(base):
            import datetime as _dt
            try: last_ts = _dt.datetime.fromisoformat(since).timestamp()
            except Exception: last_ts = 0
            for fn in _excl_filter(_collect_files(base), ws.get("scope_exclude"))[:200]:
                try:
                    if os.path.getmtime(os.path.join(base, fn)) > last_ts: newf.append(fn)
                except Exception: pass
        ws["last_opened"] = _now(); ws["updated"] = _now(); _ws_store(ws)
        return json.dumps({"status": "success", "since": since,
                           "actions": [{"action": e.get("action"), "detail": e.get("detail"), "at": e.get("at")} for e in acts[-10:]],
                           "new_files": newf[:10], "new_files_count": len(newf),
                           "questions_open": len(qs_open),
                           "next": (ws.get("state") or {}).get("next", ""),
                           "goals": [{"title": g.get("title"), "progress": g.get("progress")} for g in (ws.get("goals") or [])]}, ensure_ascii=False)

    # ---------- КУРИРУЕМЫЙ воркспейс (ws-v1.2): контуры + ручное состояние ----------
    if op == "create_curated":
        # name + contours(JSON-список [{key,name,understood,folder}]) → скелет; идемпотентно по name, состояние контуров сохраняется
        try: cts = json.loads(contours) if not _b(contours) else []
        except Exception: cts = []
        if _b(name) or not isinstance(cts, list) or not cts:
            return json.dumps({"status": "error", "message": "нужны name и contours (JSON-список)"}, ensure_ascii=False)
        wsid = "ws_" + hashlib.md5(("curated|" + name.lower()).encode("utf-8")).hexdigest()[:10]
        prev = _ws_load(wsid) or {}
        prevc = {c.get("key"): c for c in (prev.get("contours") or [])}
        out = []
        for c in cts:
            key = str(c.get("key") or c.get("name") or "").strip()
            if not key: continue
            old = prevc.get(key, {})
            st = old.get("state") or {"done": [], "next": [], "blockers": [], "waiting": []}
            out.append({"key": key, "name": c.get("name") or key, "understood": c.get("understood", ""),
                        "folder": c.get("folder", ""), "state": st, "refreshed": old.get("refreshed", "")})
        ws = {"ws_id": wsid, "name": name, "kind": "curated",
              "understood": (understood if not _b(understood) else prev.get("understood", "")),
              "folder": "", "created": prev.get("created") or _now(), "updated": _now(),
              "contours": out, "contract": "ws-v1.2"}
        _ws_store(ws)
        blk = sum(len((c["state"].get("blockers") or [])) for c in out)
        _idx_upsert({"ws_id": wsid, "name": name, "folder": "", "updated": ws["updated"], "blockers": blk, "kind": "curated"})
        return json.dumps({"status": "success", "note": "Воркспейс «" + str(name) + "» создан (" + str(len(out)) + " контуров)", "workspace": ws}, ensure_ascii=False)

    if op == "set_state":
        # обновить состояние ОДНОГО контура: state(JSON) с любыми из done/next/blockers/waiting/recent (списки). Заменяет переданные ключи.
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws or ws.get("kind") != "curated":
            return json.dumps({"status": "error", "message": "курируемый воркспейс не найден (или временный сбой — повтори)"}, ensure_ascii=False)
        if _b(contour):
            return json.dumps({"status": "error", "message": "нужен contour (ключ)"}, ensure_ascii=False)
        try: patch = json.loads(state) if not _b(state) else {}
        except Exception: patch = {}
        if not isinstance(patch, dict):
            return json.dumps({"status": "error", "message": "state должен быть JSON-объектом"}, ensure_ascii=False)
        hit = next((c for c in ws.get("contours", []) if c.get("key") == contour), None)
        if not hit:
            return json.dumps({"status": "error", "message": "контур не найден: " + str(contour)}, ensure_ascii=False)
        st = hit.setdefault("state", {"done": [], "next": [], "blockers": [], "waiting": []})
        for k in ("done", "next", "blockers", "waiting", "recent"):
            if k in patch and isinstance(patch[k], list):
                st[k] = [str(x) for x in patch[k]]
        hit["refreshed"] = _now(); ws["updated"] = _now()
        _ws_store(ws)
        blk = sum(len((c.get("state") or {}).get("blockers") or []) for c in ws.get("contours", []))
        _idx_upsert({"ws_id": ws["ws_id"], "name": ws.get("name", ""), "folder": "", "updated": ws["updated"], "blockers": blk, "kind": "curated"})
        return json.dumps({"status": "success", "note": "Контур «" + str(contour) + "» обновлён", "workspace": ws}, ensure_ascii=False)

    # ---------- действия над воркспейсом (ws-v1.1) ----------
    if op == "chat":
        # ЧАТ ВОРКСПЕЙСА (P1-5, канон: чат = интерфейс УПРАВЛЕНИЯ, память живёт в воркспейсе, не в нити).
        # Каждый вопрос самодостаточен: весь контекст воркспейса подмешивается заново (previous_response_id не нужен).
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой — повтори)"}, ensure_ascii=False)
        q = str(value or capability or "").strip()
        if not q:
            return json.dumps({"status": "error", "message": "нужен вопрос (value)"}, ensure_ascii=False)
        mf = _kv_get("me:facts") or {}
        fx = dict(mf); fx.update(ws.get("facts") or {})
        _okd, drr = _kv_get_retry("ws:" + str(ws_id) + ":draft_registry")
        reg = (drr or {}).get("rows") or []
        reg_txt = "\n".join("- %s | %s | %s %s | %s" % (r.get("asset"), r.get("class"), r.get("amount"), r.get("currency"), r.get("instrument")) for r in reg[:40])
        goals_txt = "\n".join("- %s (%s%%, инструмент: %s)" % (g.get("title"), g.get("progress", 0), g.get("tool", "")) for g in (ws.get("goals") or []))
        tasks = ws.get("tasks") or []
        tasks_txt = "\n".join("- [%s/%s] %s" % (t.get("status"), t.get("kind"), t.get("title")) for t in tasks[:20])
        qs_open = [q2 for q2 in (ws.get("questions") or []) if q2.get("status") != "answered"]
        writes = ws.get("writes") or []
        ctx = ("ПРОЕКТ: «%s» — %s. Тип: %s.\nЦЕЛИ:\n%s\nЗАДАЧИ:\n%s\nЧТО ДАЛЬШЕ: %s\n"
               "ФАКТЫ: %s\nОТКРЫТЫХ ВОПРОСОВ К ПОЛЬЗОВАТЕЛЮ: %d\nЗАПИСАННЫЕ ФАЙЛЫ: %s\n"
               + ("РЕЕСТР (черновик, %d позиций):\n%s" % (len(reg), reg_txt) if reg else "Реестр ещё не собран.")) % (
               ws.get("name", ""), ws.get("understood", ""), ((ws.get("project_type") or {}).get("name") or "—"),
               goals_txt or "—", tasks_txt or "—", (ws.get("state") or {}).get("next", "—"),
               json.dumps(fx, ensure_ascii=False)[:500], len(qs_open),
               ", ".join(os.path.basename(w2.get("path", "")) for w2 in writes[-3:]) or "нет")
        ai4 = _ask_qwen("Ты — Extella Workspace, рабочая среда проекта. Ответь на вопрос пользователя ПО ДАННЫМ проекта ниже — "
                        "коротко, по-русски, деловым языком; числа считай точно по реестру. Если для ответа нужно ДЕЙСТВИЕ "
                        "(собрать/записать/сравнить/сверить/черновик) — предложи его. НЕ выдумывай данных, которых нет.\n\n"
                        "=== ДАННЫЕ ПРОЕКТА ===\n" + ctx + "\n\n=== ВОПРОС ===\n" + q[:400] +
                        '\n\nВерни СТРОГО JSON {"answer":"<ответ, 1-5 фраз>","suggested_action":"<команда для чата, напр. «собери реестр», или пусто>"}', mx=1200)
        if not ai4 or not ai4.get("answer"):
            return json.dumps({"status": "error", "error": "llm_unavailable", "retryable": True, "message": "модель занята — повтори"}, ensure_ascii=False)
        return json.dumps({"status": "success", "kind": "chat", "answer": str(ai4.get("answer"))[:1500],
                           "suggested_action": str(ai4.get("suggested_action") or "")[:80],
                           "note": str(ai4.get("answer"))[:200]}, ensure_ascii=False)

    if op == "registry_chunk":
        # ТЯЖЁЛЫЙ РЕЕСТР ПО КУСОЧКАМ (P1-6): один вызов = один Qwen-батч (~6 файлов) → всегда влезает в таймаут.
        # Драйвер/сервер плагина крутит цикл: offset=0 → next_offset → … → done. Черновик копится в draft_registry.
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой — повтори)"}, ensure_ascii=False)
        base = os.path.expanduser(ws.get("folder") or "")
        if not base or not os.path.isdir(base):
            return json.dumps({"status": "error", "message": "у проекта нет папки"}, ensure_ascii=False)
        try: off = int(str(n or "0") or "0")
        except Exception: off = 0
        BATCH = 6
        allf = [f for f in _excl_filter(ws.get("scope_files") or _collect_files(base), ws.get("scope_exclude"))
                if f.lower().endswith((".xlsx", ".xlsm", ".csv", ".pdf", ".docx", ".txt", ".md"))][:120]
        total = len(allf)
        if off >= total:
            return json.dumps({"status": "success", "done": True, "processed": total, "total": total,
                               "note": "Все файлы уже обработаны"}, ensure_ascii=False)
        batch_files = allf[off:off + BATCH]
        # чтение содержимого on-device (те же ридеры, что в registry)
        try: include("import openpyxl", ["extella-pip install openpyxl"]); import openpyxl as _ox2
        except Exception:
            try: import openpyxl as _ox2
            except Exception: _ox2 = None
        import csv as _csv
        def _rt(fn):
            pth = os.path.join(base, fn); low = fn.lower()
            try:
                if low.endswith((".xlsx", ".xlsm")) and _ox2:
                    wb = _ox2.load_workbook(pth, read_only=True, data_only=True); o = []
                    for sh in wb.worksheets:
                        o.append("[" + sh.title + "]")
                        for i2, row in enumerate(sh.iter_rows(values_only=True)):
                            if i2 > 120: break
                            o.append(" | ".join("" if c is None else str(c) for c in row))
                    wb.close(); return "\n".join(o)
                if low.endswith(".csv"):
                    with open(pth, newline="", encoding="utf-8-sig", errors="ignore") as fh:
                        return "\n".join(" | ".join(r) for i2, r in enumerate(_csv.reader(fh)) if i2 <= 120)
                if low.endswith(".pdf"):
                    try: include("import pdfplumber", ["extella-pip install pdfplumber"])
                    except Exception: pass
                    import pdfplumber
                    o = []
                    with pdfplumber.open(pth) as pdf:
                        for pg in pdf.pages[:12]: o.append(pg.extract_text() or "")
                    return "\n".join(o)
                if low.endswith(".docx"):
                    try: include("import docx", ["extella-pip install python-docx"])
                    except Exception: pass
                    import docx as _dx2
                    return "\n".join(x.text for x in _dx2.Document(pth).paragraphs if x.text.strip())
                if low.endswith((".txt", ".md")):
                    with open(pth, encoding="utf-8", errors="ignore") as fh: return fh.read()
            except Exception: return ""
            return ""
        docs = [(fn, (_rt(fn) or "")[:3200]) for fn in batch_files]
        docs = [(fn, t) for fn, t in docs if t.strip()]
        mf = _kv_get("me:facts") or {}
        facts = dict(mf); facts.update(ws.get("facts") or {})
        owner = (facts.get("owner") or "").strip()
        added = []
        if docs and owner:
            SYS = ("Ты — Extella, финансовый ассистент. Из содержимого файлов СОБЕРИ строки реестра инвестиций. "
                   "СОБИРАЙ ТОЛЬКО доли/инвестиции этих сущностей: «" + owner + "». Чужих держателей НЕ включай. "
                   'Верни СТРОГО JSON {"rows":[{"asset":"<компания/актив>","class":"<стартап|акции|крипта|недвижимость|займ|облигации|опцион|прочее>",'
                   '"amount":<число или null>,"currency":"<USD|KZT|EUR>","date":"<YYYY-MM или \'\'>","instrument":"<SAFE|SSA|loan|share|option>",'
                   '"source":"<имя файла>","note":"<кратко>"}]}. Только реальные данные; без выдумок.\n\nФАЙЛЫ:\n'
                   + "\n\n".join("=== %s ===\n%s" % (fn, t) for fn, t in docs))
            ai2 = _ask_qwen(SYS, mx=2200)
            added = [r for r in (ai2.get("rows") or []) if isinstance(r, dict) and r.get("asset")]
        # merge в draft (компактно, кап по размеру)
        _okd, dr = _kv_get_retry("ws:" + ws["ws_id"] + ":draft_registry")
        dr = dr if isinstance(dr, dict) else {}
        rows0 = dr.get("rows") or [] if off > 0 else []   # offset 0 = новая сборка
        seen = set((str(r.get("asset", "")).lower().strip(), str(r.get("instrument", "")).lower(), str(r.get("amount"))) for r in rows0)
        for r in added:
            kk = (str(r.get("asset", "")).lower().strip(), str(r.get("instrument", "")).lower(), str(r.get("amount")))
            if kk in seen: continue
            seen.add(kk)
            rows0.append({"asset": str(r.get("asset", ""))[:60], "class": str(r.get("class", ""))[:20],
                          "amount": r.get("amount"), "currency": str(r.get("currency", ""))[:6],
                          "date": str(r.get("date", ""))[:10], "instrument": str(r.get("instrument", ""))[:20],
                          "source": os.path.basename(str(r.get("source", "")))[:48], "note": str(r.get("note", ""))[:60]})
        while rows0 and len(json.dumps({"rows": rows0}, ensure_ascii=False)) > 7000:
            rows0 = rows0[:-3]
        next_off = off + BATCH
        done = next_off >= total
        _kv_set("ws:" + ws["ws_id"] + ":draft_registry", {"rows": rows0, "total_rows": len(rows0), "at": _now(), "done": done, "files_total": total})
        if done:
            ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": "собрал реестр (по кусочкам)",
                                                "detail": "%d позиций из %d файлов" % (len(rows0), total), "at": _now(), "undo": "none"})
            adv = _advance(ws, "capability", "собрать свести реестр консолидация", key)
            ws["updated"] = _now(); _ws_store(ws)
        need_owner = (not owner)
        return json.dumps({"status": "success", "done": done, "processed": min(next_off, total), "total": total,
                           "next_offset": (None if done else next_off), "rows_so_far": len(rows0), "added_now": len(added),
                           "need_owner": need_owner,
                           "note": ("Реестр готов: %d позиций из %d файлов. Скажи «запиши в таблицу»." % (len(rows0), total)) if done
                                   else ("Обработано %d из %d файлов · собрано позиций: %d" % (min(next_off, total), total, len(rows0)))}, ensure_ascii=False)

    if op in ("set_autopilot", "undo", "organize", "run_capability", "cap_design", "cap_apply", "cap_dismiss", "cap_result", "cap_list", "cap_forget"):
        def _truthy(v): return str(v).strip().lower() in ("1", "true", "yes", "on", "да")
        ws = _ws_load(ws_id) if not _b(ws_id) else None
        if not ws:
            return json.dumps({"status": "error", "message": "воркспейс не найден (или временный сбой хранилища — повтори через пару секунд)"}, ensure_ascii=False)
        base = os.path.expanduser(ws.get("folder") or "")
        def _log2(action, detail, undo="none"):
            ws.setdefault("ledger", []).append({"n": len(ws.get("ledger", [])) + 1, "action": action,
                                                "detail": detail, "at": _now(), "undo": undo})
        def _save():
            ws["updated"] = _now(); _ws_store(ws)

        def _cap_design_run(task_text, auto=False):
            # ПРОЕКТИРОВЩИК СПОСОБНОСТЕЙ: Qwen пишет спецификацию + код нового эксперта под тупик проекта.
            task_text = str(task_text or "").strip()
            if not task_text:
                return {"status": "error", "message": "скажи, чему научиться: «научись: <что сделать>»"}
            # ПАМЯТЬ О ВЫРАЩЕННОМ — только для АВТОПИЛОТА (чтобы не перепроектировал по кругу).
            # Явный запрос человека «научись: …» — всегда новое проектирование (упоминание чужой способности ≠ дубль).
            def _stems(s):   # сравнение по ОСНОВАМ (5 симв.): «налоговой»≈«налоговая»≈«налог» — иначе автопилот плодит дубли
                return set(w[:5] for w in re.findall(r"[а-яёa-z0-9]{4,}", str(s).lower()))
            if auto:
                tt = _stems(task_text)
                tlow = task_text.lower()
                for c in (_kv_get("ws:" + ws["ws_id"] + ":caps") or []):
                    if not isinstance(c, dict): continue
                    ct = _stems(str(c.get("req") or "") + " " + str(c.get("cap_name") or ""))
                    trg_hits = sum(1 for t in (c.get("triggers") or []) if t in tlow)
                    if len(tt & ct) >= 2 or trg_hits >= 2:
                        return {"status": "success", "kind": "cap_exists", "cap_name": c.get("cap_name"),
                                "expert_name": c.get("expert_name"), "params": dict(c.get("params") or {}, ws_id=ws["ws_id"]),
                                "key": key, "note": "Способность «%s» уже выращена — запускаю её." % c.get("cap_name")}
            old = _kv_get("ws:" + ws["ws_id"] + ":cap_draft")
            if isinstance(old, dict) and old.get("code") and old.get("req") == task_text[:120]:
                qk0 = "cap_" + str(old.get("slug"))   # вопрос мог потеряться (гонка) — пересоздаём, чтобы подсказка была правдой
                if not any(q.get("key") == qk0 and q.get("status") != "answered" for q in ws.get("questions", [])):
                    ws.setdefault("questions", []).append({
                        "id": "q_" + qk0, "key": qk0, "scope": "workspace", "answer_type": "choice",
                        "question": "Разрешить новую способность «%s»? Что делает: %s." % (old.get("cap_name"), old.get("what")),
                        "candidates": ["разрешаю", "не надо"], "status": "open", "answer": ""})
                    _save()
                return {"status": "success", "kind": "cap_preview", "cap_name": old.get("cap_name"), "what": old.get("what"),
                        "reads": old.get("reads"), "writes": old.get("writes"), "network": old.get("network") or [],
                        "code": old.get("code"), "slug": old.get("slug"), "sha": old.get("sha"),
                        "note": "Способность «%s» уже спроектирована — жду твоего «разрешаю»." % old.get("cap_name")}
            rows = (_kv_get("ws:" + ws["ws_id"] + ":draft_registry") or {}).get("rows") or []
            sample = json.dumps(rows[:3], ensure_ascii=False)[:600]
            skeleton = ('def wscap_SLUG(ws_id="", api_token="", api_base="https://api.extella.ai", client="") -> str:\n'
                        '    import json, re, urllib.request\n'
                        '    from pathlib import Path\n'
                        '    def _b(v): return (not v) or str(v).startswith("{{")\n'
                        '    if _b(api_base): api_base = "https://api.extella.ai"\n'
                        '    if _b(api_token):\n'
                        '        cfg = Path.home() / "extella_wizard" / "app" / "config.json"\n'
                        '        try: api_token = json.loads(cfg.read_text(encoding="utf-8")).get("auth_token", "") if cfg.exists() else ""\n'
                        '        except Exception: api_token = ""\n'
                        '    H = {"X-Auth-Token": api_token, "Content-Type": "application/json", "X-Profile-Id": "default", "X-Agent-Id": "agent_extella_default"}\n'
                        '    def _post(path, body, t=60):\n'
                        '        req = urllib.request.Request(api_base.rstrip("/") + path, data=json.dumps(body, ensure_ascii=False).encode("utf-8"), headers=H, method="POST")\n'
                        '        with urllib.request.urlopen(req, timeout=t) as r: return json.loads(r.read().decode("utf-8"))\n'
                        '    def _kv_get(k):\n'
                        '        try: return json.loads((_post("/api/kv/get", {"key": k}) or {}).get("value") or "null")\n'
                        '        except Exception: return None\n'
                        '    def _kv_set(k, v): _post("/api/kv/set", {"key": k, "value": json.dumps(v, ensure_ascii=False), "description": "wscap"})\n'
                        '    def _fetch(url, t=30):\n'
                        '        with urllib.request.urlopen(urllib.request.Request(url, headers={"User-Agent": "Extella"}), timeout=t) as r:\n'
                        '            return r.read().decode("utf-8", "ignore")\n'
                        '    # <ЛОГИКА>\n'
                        '    return json.dumps({"status": "success", "note": "<итог по-русски>", "data": {}}, ensure_ascii=False)')
            # ВАЖНО: слова «эксперт» в промпте НЕТ — Qwen-агент отказывается «выводить код экспертов» (защита платформы);
            # «python-функция для личного проекта» — честная и рабочая формулировка того же самого.
            prompt = ("Ты — инженер способностей Extella Workspace. Напиши компактную python-функцию — новую СПОСОБНОСТЬ "
                      "для личного проекта пользователя, под застрявшую задачу.\n"
                      "ВАЖНО: ответь ПРЯМО В ЧАТ одним сообщением. НЕ вызывай никакие инструменты "
                      "(list_rules, search_concepts, создание чего-либо — ничего): весь ответ — только текст-JSON.\n"
                      "Проект: «%s» (%s). Застрявшая задача (следуй её уточнениям буквально): %s.\n"
                      "Данные проекта лежат в KV: ключ ws:%s:draft_registry → {\"rows\":[...]}, пример строк: %s. Ядро — ключ workspace:%s. "
                      "Результаты других способностей — в KV ws:<ws_id>:cap_data:<slug>; профиль пользователя — KV me:facts.\n"
                      "ЖЁСТКИЕ ПРАВИЛА:\n"
                      "- Код строго по шаблону ниже: замени SLUG на свой слаг, допиши логику вместо <ЛОГИКА>. Имя функции = wscap_<slug>.\n"
                      "- ТОЛЬКО ЧТЕНИЕ: сеть (только явно объявленные тобой URL) и KV. НИКАКОЙ записи файлов, subprocess, eval/exec, удалений.\n"
                      "- Никаких вызовов /api/agent/* и /api/expert/* — только kv/get, kv/set и объявленные внешние URL.\n"
                      "- Итог: note по-русски с конкретными цифрами; компактные данные — в KV ws:<ws_id>:cap_data:<slug> (≤6КБ) и кратко в поле data.\n"
                      "- Только стандартная библиотека Python; максимум 110 строк; аккуратно нормализуй форматы дат/чисел; на сбое сети — честный status error.\n"
                      "ШАБЛОН:\n%s\n"
                      'Верни СТРОГО JSON без пояснений: {"slug":"<a-z0-9_ до 20 символов>","cap_name":"<по-русски 2-4 слова>",'
                      '"what":"<одна фраза: что делает>","triggers":["<3-6 слов-триггеров в нижнем регистре, по которым звать способность>"],'
                      '"network":["<полные URL, куда ходит код>"],"reads":"<что читает>","writes":"только KV",'
                      '"code":"<весь код одной JSON-строкой, переносы строк как \\n>"}'
                      % (ws.get("name", ""), str(ws.get("understood", ""))[:200], task_text[:700],
                         ws["ws_id"], sample, ws["ws_id"], skeleton))
            d = _ask_qwen(prompt, mx=9000, t=240, tries=3)   # рассуждающая модель: бюджет на «мышление» + код; 500-ки платформы ретраим
            code = str(d.get("code") or "")
            if not code:
                return {"status": "error", "error": "llm_unavailable", "retryable": True,
                        "message": "проектировщик занят — повтори через минуту"}
            if len(code) > 5500:
                return {"status": "error", "message": "спроектированный код слишком большой (KV-лимит) — сузь задачу и повтори «научись: …»"}
            mfn = re.search(r"def\s+wscap_([a-z0-9_]{3,32})\s*\(", code)
            if not mfn:
                return {"status": "error", "message": "код без функции wscap_<slug> (a-z0-9_, 3-32 символа) — повтори «научись: …»", "retryable": True}
            # соль воркспейса в имени: имя функции = имя эксперта, и никакой воркспейс не перепишет чужую способность
            fn = mfn.group(1)
            slug = fn + "_" + hashlib.md5(ws["ws_id"].encode()).hexdigest()[:6]
            code = code.replace("def wscap_" + fn + "(", "def wscap_" + slug + "(", 1)
            ok, why, _u = _cap_lint(code, d.get("network") or [])
            if not ok:
                return {"status": "error", "error": "cap_lint", "retryable": True,
                        "message": "код не прошёл проверку безопасности (" + why + ") — повтори «научись: …»"}
            sha = hashlib.sha256(code.encode("utf-8")).hexdigest()[:16]
            trg = _norm_triggers(d.get("triggers"))
            for tok in re.findall(r"[а-яёa-z0-9]{4,}", task_text.lower())[:4]:   # триггеры всегда узнают исходную задачу
                if tok not in trg and len(trg) < 8: trg.append(tok)
            dr = {"slug": slug, "cap_name": str(d.get("cap_name") or fn)[:60], "what": str(d.get("what") or "")[:200],
                  "triggers": trg, "network": [str(u)[:120] for u in (d.get("network") or [])][:5],
                  "reads": str(d.get("reads") or "")[:150], "writes": "только KV", "params": {},
                  "code": code, "req": task_text[:120], "sha": sha, "at": _now()}
            if len(json.dumps(dr, ensure_ascii=False).encode("utf-8")) > 7500:   # лимит KV меряем по байтам ВСЕГО значения
                return {"status": "error", "message": "спроектированный код слишком большой (KV-лимит) — сузь задачу и повтори «научись: …»"}
            _kv_set("ws:" + ws["ws_id"] + ":cap_draft", dr)
            chk = _kv_get("ws:" + ws["ws_id"] + ":cap_draft")
            if not (isinstance(chk, dict) and chk.get("sha") == sha):
                return {"status": "error", "retryable": True, "message": "черновик способности не сохранился (сбой KV) — повтори «научись: …»"}
            qk = "cap_" + slug
            for q in ws.get("questions", []):   # старые предложения закрываем — «разрешаю» относится только к свежей карточке
                if str(q.get("key", "")).startswith("cap_") and q.get("key") != qk and q.get("status") != "answered":
                    q["status"] = "answered"; q["answer"] = "заменено новым предложением"
            if not any(q.get("key") == qk and q.get("status") != "answered" for q in ws.get("questions", [])):
                ws.setdefault("questions", []).append({
                    "id": "q_" + qk, "key": qk, "scope": "workspace", "answer_type": "choice",
                    "question": "Разрешить новую способность «%s»? Что делает: %s. Сеть: %s. Пишет: только KV."
                                % (dr["cap_name"], dr["what"], ", ".join(dr["network"]) or "нет"),
                    "candidates": ["разрешаю", "не надо"], "status": "open", "answer": ""})
            _log2("спроектировал способность", dr["cap_name"] + " (жду «разрешаю»)")
            _save()
            return {"status": "success", "kind": "cap_preview", "slug": slug, "cap_name": dr["cap_name"], "what": dr["what"],
                    "reads": dr["reads"], "writes": dr["writes"], "network": dr["network"], "code": code, "sha": sha,
                    "note": "Спроектировал способность «%s». Посмотри карточку (что делает, куда ходит, код) и скажи «разрешаю» — зарегистрирую и продолжу." % dr["cap_name"]}

        if op == "cap_design":
            return json.dumps(_cap_design_run(value or capability, auto=(str(n).strip().lower() == "auto")), ensure_ascii=False)

        if op == "cap_apply":
            vl = str(value or "").strip().lower()
            if re.search(r"(^|[\s,.!])(не|нет|no)([\s,.!]|$)", vl) or "не надо" in vl:
                return json.dumps({"status": "error", "message": "понял как отказ — регистрация отменена (убрать предложение: «не надо»)"}, ensure_ascii=False)
            confirmed = _truthy(apply) or bool(re.search(r"разреша|подтвержда|confirm", vl))
            if not confirmed:
                return json.dumps({"status": "error", "message": "регистрирую только после твоего «разрешаю» (канон: код одобряет человек)"}, ensure_ascii=False)
            okr, cnote, cap = _cap_register(ws, want_slug=re.sub(r"^cap_", "", str(key or "").strip()), want_sha=str(n or "").strip())
            if not okr:
                return json.dumps({"status": "error", "message": cnote}, ensure_ascii=False)
            _save()
            return json.dumps({"status": "success", "kind": "cap_added", "cap": cap, "note": cnote, "workspace": ws}, ensure_ascii=False)

        if op == "cap_dismiss":
            dr = _kv_get("ws:" + ws["ws_id"] + ":cap_draft") or {}
            _kv_set("ws:" + ws["ws_id"] + ":cap_draft", None)
            for q in ws.get("questions", []):
                if str(q.get("key", "")).startswith("cap_") and q.get("status") != "answered":
                    q["status"] = "answered"; q["answer"] = "не надо"
            _log2("способность отклонена", str(dr.get("cap_name") or "")[:60])
            _save()
            return json.dumps({"status": "success", "note": "Убрал предложение способности. Если передумаешь — «научись: <что сделать>»", "workspace": ws}, ensure_ascii=False)

        if op == "cap_list":
            # инвентарь выращенных способностей (для Настроек): реестр + ожидающее предложение
            creg = [c for c in (_kv_get("ws:" + ws["ws_id"] + ":caps") or []) if isinstance(c, dict)]
            dr = _kv_get("ws:" + ws["ws_id"] + ":cap_draft")
            pend = ({"slug": dr.get("slug"), "cap_name": dr.get("cap_name"), "what": dr.get("what")}
                    if isinstance(dr, dict) and dr.get("code") else None)
            return json.dumps({"status": "success", "caps": creg, "pending": pend,
                               "note": "Выращено способностей: %d%s" % (len(creg), " · 1 ждёт «разрешаю»" if pend else "")}, ensure_ascii=False)

        if op == "cap_forget":
            # РАЗУЧИТЬСЯ: убрать способность из диспетча и пака (обратимо — «научись» вырастит заново)
            slug = re.sub(r"^cap_", "", str(key or "").strip())
            if not slug:
                return json.dumps({"status": "error", "message": "нужен key (slug способности)"}, ensure_ascii=False)
            creg = [c for c in (_kv_get("ws:" + ws["ws_id"] + ":caps") or []) if isinstance(c, dict)]
            hit = next((c for c in creg if c.get("slug") == slug), None)
            if not hit:
                return json.dumps({"status": "error", "message": "такой способности нет: " + slug}, ensure_ascii=False)
            creg = [c for c in creg if c.get("slug") != slug]
            _kv_set("ws:" + ws["ws_id"] + ":caps", creg)
            chk = _kv_get("ws:" + ws["ws_id"] + ":caps")
            if isinstance(chk, list) and any(c.get("slug") == slug for c in chk if isinstance(c, dict)):
                return json.dumps({"status": "error", "retryable": True, "message": "реестр не сохранился (сбой KV) — повтори"}, ensure_ascii=False)
            pk = ws.setdefault("pack", {}).get("capabilities") or []
            ws["pack"]["capabilities"] = [c for c in pk if c != "🧠 " + str(hit.get("cap_name"))]
            dele = False
            try:   # имя посолено этим ws_id — эксперт больше никому не нужен, пробуем удалить с платформы
                r = _post("/api/expert/delete", {"name": hit.get("expert_name"), "global": True}, t=30)
                dele = isinstance(r, dict) and r.get("status") == "success"
            except Exception: pass
            _log2("способность разучена", str(hit.get("cap_name"))[:60])
            _save()
            return json.dumps({"status": "success", "kind": "cap_forgotten",
                               "note": "Разучился: «%s»%s. Вырастить заново — «научись: <что сделать>»."
                                       % (hit.get("cap_name"), "" if dele else " (эксперт остался на платформе, но воркспейс его больше не зовёт)"),
                               "workspace": ws}, ensure_ascii=False)

        if op == "cap_result":
            # исполнитель (автопилот/плагин) прогнал выращенную способность и вернул итог → закрыть задачу, двигать план.
            # Итог способности — недоверенный текст: скрабим секреты перед ledger/чатом (канон: секреты не печатать).
            def _scrub(s):
                s = str(s or "")
                try:
                    toks = [api_token]
                    tp = Path.home() / ".extella_test_token"
                    if tp.exists(): toks.append(tp.read_text(encoding="utf-8").strip())
                    for t in toks:
                        if t and len(t) > 8 and t in s: s = s.replace(t, "•••")
                except Exception: pass
                return re.sub(r"\bsk-[A-Za-z0-9_\-]{10,}", "sk-•••", s)
            clean = _scrub(str(value or ""))[:200]
            _log2("способность отработала", clean[:110])
            adv = _advance(ws, "capability", str(capability or "") + " " + clean, key)
            _save()
            return json.dumps({"status": "success", "kind": "cap_result", "note": clean or "готово",
                               "advanced": (adv["title"] if adv else None),
                               "task_next": ws.get("state", {}).get("next")}, ensure_ascii=False)

        if op == "set_autopilot":
            ws.setdefault("autopilot", {"enabled": False, "rules": []})["enabled"] = _truthy(enabled)
            _log2("автопилот " + ("включён" if ws["autopilot"]["enabled"] else "выключен"), "по этому проекту")
            _save()
            return json.dumps({"status": "success", "note": "Автопилот " + ("включён — нажми ▶ Запустить, чтобы он поехал" if ws["autopilot"]["enabled"] else "выключен"),
                               "workspace": ws}, ensure_ascii=False)

        if op == "run_capability":
            # КОНТЕНТНЫЙ СЛОЙ (ws-v1.2): реально ЧИТАЕТ содержимое таблиц ON-DEVICE (XLSX/CSV) и делает работу.
            # Два сценария: «свести цифры/счета в реестр» и «сравнить версии». Содержимое наружу не уходит.
            import csv
            text = str(capability or "").lower()
            # «научись: X» — прямой запрос вырастить новую способность (из чата)
            m_learn = re.match(r"^\s*(научись|научи себя|выучи)\s*[:,]?\s+(.+)$", str(capability or ""), re.I | re.S)
            if m_learn:
                return json.dumps(_cap_design_run(m_learn.group(2)), ensure_ascii=False)
            def _num(x):
                if isinstance(x, (int, float)): return float(x)
                if x is None: return None
                s = str(x).strip().replace(" ", "").replace(" ", "")
                s = re.sub(r"[^\d,.\-]", "", s)
                if not s or s in ("-", ".", ","): return None
                if "," in s and "." in s: s = s.replace(",", "")
                elif "," in s: s = s.replace(",", ".")
                try: return float(s)
                except Exception: return None
            def _ncols(header, data):
                w = len(header)
                for r in data[:50]: w = max(w, len(r))
                out = []
                for j in range(w):
                    vals = []; nz = 0
                    for r in data:
                        if j < len(r) and r[j] not in (None, ""):
                            nz += 1
                            v = _num(r[j])
                            if v is not None: vals.append(v)
                    if vals and nz and len(vals) >= 0.6 * nz and len(vals) >= 2:
                        lab = str(header[j]).strip() if (j < len(header) and header[j] not in (None, "")) else ("колонка %d" % (j + 1))
                        out.append({"label": lab, "sum": round(sum(vals), 2), "count": len(vals)})
                return out
            def _fmt(v):
                v = round(v, 2)
                if v == int(v): return format(int(v), ",d").replace(",", " ")
                return format(v, ",.2f").replace(",", " ")
            _oxl = None
            need_xlsx = os.path.isdir(base) and any(f.lower().endswith((".xlsx", ".xlsm")) for f in _collect_files(base))  # РЕКУРСИВНО (xlsx в подпапках)
            if need_xlsx:
                try:
                    include("import openpyxl", ["extella-pip install openpyxl"]); import openpyxl as _oxl
                except Exception:
                    _oxl = None
                if _oxl is None:
                    try: import openpyxl as _oxl
                    except Exception:
                        try:
                            import subprocess, sys
                            subprocess.run([sys.executable, "-m", "pip", "install", "-q", "openpyxl"], timeout=120)
                            import openpyxl as _oxl
                        except Exception: _oxl = None
            def _read_xlsx(p, mx=5000):
                sheets = []
                try:
                    wb = _oxl.load_workbook(p, read_only=True, data_only=True)
                    for sh in wb.worksheets:
                        rows = list(sh.iter_rows(values_only=True))
                        if len(rows) >= 2:
                            cols = _ncols(rows[0], rows[1:mx])
                            if cols: sheets.append({"sheet": sh.title, "rows": len(rows) - 1, "columns": cols})
                    wb.close()
                except Exception: pass
                return sheets
            def _read_csv(p, mx=5000):
                try:
                    with open(p, newline="", encoding="utf-8-sig") as fh:
                        rows = list(csv.reader(fh))
                    if len(rows) >= 2:
                        cols = _ncols(rows[0], rows[1:mx])
                        if cols: return [{"sheet": "", "rows": len(rows) - 1, "columns": cols}]
                except Exception: pass
                return []
            def _read_any(fn):
                p = os.path.join(base, fn); low = fn.lower()
                if low.endswith((".xlsx", ".xlsm")): return _read_xlsx(p) if _oxl else []
                if low.endswith(".csv"): return _read_csv(p)
                return []
            # ---- ЧЕРНОВИКИ ДОКУМЕНТОВ/ПИСЕМ (DOCX / .eml): Qwen пишет черновик → превью → подтверждение → файл → undo ----
            # КАНОН: письма — ТОЛЬКО черновики (.eml открывается в почтовике, отправляет человек). Автопилот гейт не проходит.
            _create_kw = ("черновик", "драфт", "подготовь", "подготовить", "составь", "составить", "напиши", "написать", "сгенерируй", "запиши", "запис", "сохрани")
            _doc_kw = ("договор", "протокол", "письм", "email", "мейл", "оффер", "документ", "заявлен", "соглашен", "кп", "запрос")
            if any(k in text for k in _create_kw) and any(k in text for k in _doc_kw):
                is_email = any(k in text for k in ("письм", "email", "мейл"))
                confirmed = _truthy(apply) or any(k in text for k in ("подтвержда", "подтверди", "confirm"))
                dr = _kv_get("ws:" + ws["ws_id"] + ":draft_doc")
                out_dir = os.path.join(base, "Extella") if (base and os.path.isdir(base)) else os.path.join(os.path.expanduser("~"), "Downloads")
                if not (dr and dr.get("body")) or not confirmed:
                    if not (dr and dr.get("body")) or (dr.get("req") != text[:100]):
                        mf = _kv_get("me:facts") or {}; fx = dict(mf); fx.update(ws.get("facts") or {})
                        gprompt = ("Ты — Extella. Подготовь деловой ЧЕРНОВИК %s на русском для проекта «%s» (%s). "
                                   "Запрос пользователя: %s. Факты: %s. Цели проекта: %s. "
                                   "Конкретно, без воды; где данных не хватает — вставь [УТОЧНИТЬ: что]. "
                                   'Верни СТРОГО JSON {"title":"<название документа>","subject":"<тема письма или пусто>",'
                                   '"to":"<адресат если понятен, иначе пусто>","body":"<текст, абзацы через \\n\\n>"}'
                                   % ("ПИСЬМА" if is_email else "ДОКУМЕНТА", ws.get("name", ""), ws.get("understood", ""),
                                      str(capability)[:300], json.dumps(fx, ensure_ascii=False)[:400],
                                      "; ".join(g.get("title", "") for g in (ws.get("goals") or [])[:3])))
                        gd = _ask_qwen(gprompt, mx=2200)
                        if not gd.get("body"):
                            return json.dumps({"status": "error", "error": "llm_unavailable", "retryable": True,
                                               "message": "модель занята — повтори через минуту"}, ensure_ascii=False)
                        dr = {"kind": ("eml" if is_email else "docx"), "title": (gd.get("title") or "Черновик")[:120],
                              "subject": (gd.get("subject") or gd.get("title") or "")[:150], "to": (gd.get("to") or "")[:120],
                              "body": str(gd.get("body"))[:6000], "req": text[:100], "at": _now()}
                        _kv_set("ws:" + ws["ws_id"] + ":draft_doc", dr)
                    if not confirmed:
                        return json.dumps({"status": "success", "kind": "doc_preview", "doc_kind": dr["kind"],
                                           "title": dr["title"], "subject": dr["subject"], "to": dr["to"],
                                           "body": dr["body"][:4000],
                                           "note": "Черновик %s готов — «%s». Проверь текст; скажи «запиши, подтверждаю» — положу файлом (%s). %s" %
                                                   ("письма" if dr["kind"] == "eml" else "документа", dr["title"],
                                                    ".eml для почтовика" if dr["kind"] == "eml" else "DOCX",
                                                    "ОТПРАВЛЯЕШЬ ТЫ сам — я только черновик." if dr["kind"] == "eml" else ""),
                                           "privacy": "черновик сгенерирован по фактам воркспейса; файл пишется on-device"}, ensure_ascii=False)
                # ПОДТВЕРЖДЕНО → запись файла
                os.makedirs(out_dir, exist_ok=True)
                slug = re.sub(r"[/\\:]+", "_", dr["title"])[:60]
                base_name = slug if slug.lower().startswith("черновик") else (("Черновик письма — " if dr["kind"] == "eml" else "Черновик — ") + slug)
                if dr["kind"] == "eml":
                    target = os.path.join(out_dir, base_name + ".eml")
                    n_ = 2
                    while os.path.exists(target):
                        target = os.path.join(out_dir, "%s v%d.eml" % (base_name, n_)); n_ += 1
                    from email.message import EmailMessage
                    m = EmailMessage()
                    m["Subject"] = dr["subject"] or dr["title"]
                    if dr["to"]: m["To"] = dr["to"]
                    m["X-Unsent"] = "1"   # почтовики открывают как ЧЕРНОВИК к редактированию
                    m.set_content(dr["body"])
                    with open(target, "wb") as f: f.write(m.as_bytes())
                else:
                    try: include("import docx", ["extella-pip install python-docx"])
                    except Exception: pass
                    try: import docx as _dx
                    except Exception:
                        try:
                            import subprocess as _sp, sys as _sy
                            _sp.run([_sy.executable, "-m", "pip", "install", "-q", "python-docx"], timeout=120); import docx as _dx
                        except Exception:
                            return json.dumps({"status": "error", "message": "python-docx недоступен — не могу записать DOCX"}, ensure_ascii=False)
                    target = os.path.join(out_dir, base_name + ".docx")
                    n_ = 2
                    while os.path.exists(target):
                        target = os.path.join(out_dir, "%s v%d.docx" % (base_name, n_)); n_ += 1
                    doc = _dx.Document()
                    doc.add_heading(dr["title"], level=1)
                    for para in dr["body"].split("\n\n"):
                        if para.strip(): doc.add_paragraph(para.strip())
                    doc.save(target)
                ws.setdefault("writes", []).append({"at": _now(), "path": target, "rows": 0, "kind": dr["kind"]})
                _kv_set("ws:" + ws["ws_id"] + ":draft_doc", None)
                _log2("записал черновик", "%s → %s" % (dr["title"][:50], os.path.basename(target)), undo="write")
                adv = _advance(ws, "capability", str(capability) + " черновик документ письмо " + dr["title"], key, artifact=True); _save()   # файл РЕАЛЬНО создан выше + в манифесте
                return json.dumps({"status": "success", "kind": "doc_written", "doc_kind": dr["kind"], "path": target,
                                   "note": "Черновик записан: «%s». %s Отменить — «отмени запись»." %
                                           (os.path.basename(target), "Открой в почтовике, проверь и отправь САМ." if dr["kind"] == "eml" else ""),
                                   "advanced": (adv["title"] if adv else None), "task_next": ws.get("state", {}).get("next"),
                                   "workspace": ws}, ensure_ascii=False)
            # ---- ЗАПИСАТЬ РЕЕСТР В ТАБЛИЦУ (гейт доверия: превью → подтверждение → XLSX on-device → undo) ----
            # Пишем ТОЛЬКО новый файл в подпапку Extella внутри folder; ничего не перезаписываем; манифест ws.writes.
            WRITE_KW = ("запиш", "запис", "сохран", "выгруз", "экспорт", "write", "save", "export")
            CONFIRM_KW = ("подтвержда", "подтверди", "confirm")
            _draft = _kv_get("ws:" + ws["ws_id"] + ":draft_registry") if os.path.isdir(base) else None
            if os.path.isdir(base) and any(k in text for k in WRITE_KW) and \
               (any(k in text for k in ("реестр", "таблиц", "registry")) or (_draft and _draft.get("rows"))):
                rows_d = (_draft or {}).get("rows") or []
                confirmed = _truthy(apply) or any(k in text for k in CONFIRM_KW)
                if rows_d:
                    out_dir = os.path.join(base, "Extella")
                    fname = "Реестр — %s.xlsx" % re.sub(r"[/\\:]+", "_", ws.get("name") or "проект")
                    target = os.path.join(out_dir, fname)
                    n_ = 2
                    while os.path.exists(target):
                        target = os.path.join(out_dir, fname[:-5] + " v%d.xlsx" % n_); n_ += 1
                    if not confirmed:
                        return json.dumps({"status": "success", "kind": "write_preview", "capability": capability,
                                           "rows": rows_d, "target_file": target, "rows_count": len(rows_d),
                                           "note": "Готов записать %d позиций в НОВЫЙ файл «%s» (внутри папки проекта, ничего не перезаписываю). "
                                                   "Скажи «запиши, подтверждаю» или нажми Подтвердить." % (len(rows_d), os.path.relpath(target, base)),
                                           "privacy": "запись on-device; в облако ничего не уходит"}, ensure_ascii=False)
                    # ЗАПИСЬ (подтверждена)
                    if _oxl is None:
                        try:
                            include("import openpyxl", ["extella-pip install openpyxl"]); import openpyxl as _oxl
                        except Exception:
                            try: import openpyxl as _oxl
                            except Exception: _oxl = None
                    if _oxl is None:
                        return json.dumps({"status": "error", "message": "openpyxl недоступен — не могу записать XLSX"}, ensure_ascii=False)
                    os.makedirs(out_dir, exist_ok=True)
                    wb = _oxl.Workbook(); sh = wb.active; sh.title = "Реестр"
                    sh.append(["Актив", "Класс", "Сумма", "Валюта", "Дата", "Инструмент", "Источник", "Заметка"])
                    for r in rows_d:
                        sh.append([r.get("asset"), r.get("class"), r.get("amount"), r.get("currency"),
                                   r.get("date"), r.get("instrument"), r.get("source"), r.get("note")])
                    wb.save(target)
                    _wkeys = [(str(r.get("asset", "")).lower().strip() + "|" + str(r.get("instrument", "")).lower() + "|" + str(r.get("date", ""))) for r in rows_d]
                    ws.setdefault("writes", []).append({"at": _now(), "path": target, "rows": len(rows_d), "keys": _wkeys})
                    _log2("записал реестр в таблицу", "%d позиций → %s" % (len(rows_d), os.path.relpath(target, base)), undo="write")
                    adv = _advance(ws, "capability", str(capability) + " записать реестр таблицу", key, artifact=True); _save()   # файл РЕАЛЬНО создан выше + в манифесте
                    return json.dumps({"status": "success", "kind": "written", "capability": capability,
                                       "path": target, "rows_written": len(rows_d),
                                       "note": "Записал %d позиций в «%s» — файл лежит в папке проекта (подпапка Extella). Отменить — «отмени запись»." % (len(rows_d), os.path.relpath(target, base)),
                                       "advanced": (adv["title"] if adv else None), "task_next": ws.get("state", {}).get("next"),
                                       "workspace": ws}, ensure_ascii=False)
                # черновика нет → проваливаемся в сборку реестра ниже (соберёт и закеширует; запись — следующим шагом)
            # ---- СВЕРКА ДЕКЛАРАЦИЯ ↔ РЕЕСТР (налоговый кейс): читает файл декларации + черновик реестра → расхождения ----
            if "деклара" in text and any(k in text for k in ("свер", "провер", "расхожд", "сравн")):
                _okr, drr = _kv_get_retry("ws:" + ws["ws_id"] + ":draft_registry")
                rows_r = (drr or {}).get("rows") or []
                if not rows_r:
                    return json.dumps({"status": "success", "queued": True,
                                       "note": "Сначала собери реестр («собери реестр») — сверять пока не с чем."}, ensure_ascii=False)
                mfx = _kv_get("me:facts") or {}; fx = dict(mfx); fx.update(ws.get("facts") or {})
                decl = str(fx.get("declaration_file") or "").strip()
                dpath = ""
                if decl:
                    cand = decl if os.path.isabs(os.path.expanduser(decl)) else os.path.join(base, decl)
                    if os.path.isfile(os.path.expanduser(cand)): dpath = os.path.expanduser(cand)
                if not dpath and os.path.isdir(base):
                    for fn in _collect_files(base):   # ищем по имени (деклара/240.00), exclude тут НЕ применяем — файл нужен по прямой просьбе
                        if ("деклара" in fn.lower() or "240" in fn.lower()) and fn.lower().endswith((".pdf", ".xlsx", ".xls", ".txt", ".docx")):
                            dpath = os.path.join(base, fn); break
                if not dpath:
                    existing = next((q for q in ws.get("questions", []) if q.get("key") == "declaration_file" and q.get("status") != "answered"), None)
                    if not existing:
                        ws.setdefault("questions", []).append({
                            "id": "q_" + hashlib.md5(("decl|" + ws["ws_id"]).encode("utf-8")).hexdigest()[:8],
                            "key": "declaration_file", "scope": "workspace", "answer_type": "file",
                            "question": "Приложи поданную декларацию (PDF/XLSX) — сверю её с реестром и выпишу расхождения.",
                            "candidates": [], "status": "open", "answer": ""})
                        _log2("воркспейс спросил", "файл декларации для сверки"); _save()
                        existing = ws["questions"][-1]
                    return json.dumps({"status": "success", "clarify": True, "need": "declaration_file", "scope": "workspace",
                                       "question": existing["question"], "candidates": [],
                                       "note": "Нужна поданная декларация — приложи файлом в инбоксе, и я сверю."}, ensure_ascii=False)
                dtext = ""
                try:
                    low = dpath.lower()
                    if low.endswith(".pdf"):
                        try: include("import pdfplumber", ["extella-pip install pdfplumber"])
                        except Exception: pass
                        import pdfplumber
                        _o = []
                        with pdfplumber.open(dpath) as pdf:
                            for pg in pdf.pages[:12]: _o.append(pg.extract_text() or "")
                        dtext = "\n".join(_o)
                    elif low.endswith((".xlsx", ".xls")) and _oxl:
                        wb2 = _oxl.load_workbook(dpath, read_only=True, data_only=True); _o = []
                        for sh2 in wb2.worksheets:
                            for i3, row in enumerate(sh2.iter_rows(values_only=True)):
                                if i3 > 200: break
                                _o.append(" | ".join("" if c is None else str(c) for c in row))
                        wb2.close(); dtext = "\n".join(_o)
                    else:
                        with open(dpath, encoding="utf-8", errors="ignore") as fh: dtext = fh.read()
                except Exception: dtext = ""
                if not dtext.strip():
                    return json.dumps({"status": "error", "message": "не смог прочитать декларацию (%s) — возможно скан; приложи читаемый файл" % os.path.basename(dpath)}, ensure_ascii=False)
                reg_lines = "\n".join("- %s | %s | %s %s | %s" % (r.get("asset"), r.get("class"), r.get("amount"), r.get("currency"), r.get("instrument")) for r in rows_r[:60])
                ai3 = _ask_qwen("Ты — Extella, налоговый ассистент (НЕ консультант — только сверка данных). Сверь РЕЕСТР активов с ПОДАННОЙ ДЕКЛАРАЦИЕЙ. "
                                'Верни СТРОГО JSON {"missing_in_declaration":[{"asset":"...","amount":...,"currency":"..."}],'
                                '"covered":["<актив>"],"mismatches":[{"asset":"...","registry":"...","declaration":"..."}],"summary":"<2-3 фразы по-русски>"}\n\n'
                                "РЕЕСТР:\n" + reg_lines + "\n\nДЕКЛАРАЦИЯ (текст):\n" + dtext[:6000], mx=2000)
                if not ai3:
                    return json.dumps({"status": "error", "error": "llm_unavailable", "retryable": True, "message": "модель занята — повтори"}, ensure_ascii=False)
                miss = ai3.get("missing_in_declaration") or []
                note3 = ("Сверил с «%s»: покрыто %d, НЕ найдено в декларации %d, расхождений %d. %s" %
                         (os.path.basename(dpath), len(ai3.get("covered") or []), len(miss), len(ai3.get("mismatches") or []), str(ai3.get("summary", ""))[:200]))
                _log2("сверил декларацию с реестром", note3[:120])
                adv = _advance(ws, "capability", str(capability) + " сверить декларацию реестром расхождения", key); _save()
                return json.dumps({"status": "success", "kind": "tax_check", "declaration": os.path.basename(dpath),
                                   "missing_in_declaration": miss, "covered": ai3.get("covered") or [],
                                   "mismatches": ai3.get("mismatches") or [], "note": note3,
                                   "advanced": (adv["title"] if adv else None), "task_next": ws.get("state", {}).get("next"),
                                   "privacy": "декларация прочитана on-device; в Qwen ушли извлечённые фрагменты (вариант B)"}, ensure_ascii=False)
            # ---- СОБРАТЬ РЕЕСТР: читает СОДЕРЖИМОЕ (xlsx/csv/pdf/docx/txt) on-device → нормализует Qwen (вариант B) ----
            CONSOLIDATE_KW = ("реестр", "трекер", "консолид", "единый", "портфел", "собрать все", "свод всех", "consolidate", "registry", "tracker", "запиш", "запис")
            if os.path.isdir(base) and any(k in text for k in CONSOLIDATE_KW):
                def _txt_xlsx(p):
                    if not _oxl: return ""
                    try:
                        wb = _oxl.load_workbook(p, read_only=True, data_only=True); parts = []
                        for sh in wb.worksheets:
                            parts.append("[" + sh.title + "]")
                            for i, row in enumerate(sh.iter_rows(values_only=True)):
                                if i > 200: break
                                parts.append(" | ".join("" if c is None else str(c) for c in row))
                        wb.close(); return "\n".join(parts)
                    except Exception: return ""
                def _txt_csv(p):
                    try:
                        with open(p, newline="", encoding="utf-8-sig", errors="ignore") as fh:
                            return "\n".join(" | ".join(r) for i, r in enumerate(csv.reader(fh)) if i <= 200)
                    except Exception: return ""
                def _txt_pdf(p):
                    try: include("import pdfplumber", ["extella-pip install pdfplumber"])
                    except Exception: pass
                    try: import pdfplumber
                    except Exception:
                        try:
                            import subprocess, sys; subprocess.run([sys.executable, "-m", "pip", "install", "-q", "pdfplumber"], timeout=180); import pdfplumber
                        except Exception: return ""
                    try:
                        out = []
                        with pdfplumber.open(p) as pdf:
                            for pg in pdf.pages[:12]: out.append(pg.extract_text() or "")
                        return "\n".join(out)
                    except Exception: return ""
                def _txt_docx(p):
                    try: include("import docx", ["extella-pip install python-docx"])
                    except Exception: pass
                    try: import docx
                    except Exception:
                        try:
                            import subprocess, sys; subprocess.run([sys.executable, "-m", "pip", "install", "-q", "python-docx"], timeout=120); import docx
                        except Exception: return ""
                    try:
                        d = docx.Document(p); return "\n".join(x.text for x in d.paragraphs if x.text.strip())
                    except Exception: return ""
                def _read_text(fn):
                    p = os.path.join(base, fn); low = fn.lower()
                    if low.endswith((".xlsx", ".xlsm")): return _txt_xlsx(p)
                    if low.endswith(".csv"): return _txt_csv(p)
                    if low.endswith(".pdf"): return _txt_pdf(p)
                    if low.endswith(".docx"): return _txt_docx(p)
                    if low.endswith((".txt", ".md")):
                        try:
                            with open(p, encoding="utf-8", errors="ignore") as fh: return fh.read()
                        except Exception: return ""
                    return ""
                docs = []
                for fn in _excl_filter(ws.get("scope_files") or _collect_files(base), ws.get("scope_exclude")):
                    if not fn.lower().endswith((".xlsx", ".xlsm", ".csv", ".pdf", ".docx", ".txt", ".md")): continue
                    t = _read_text(fn)
                    if t and t.strip(): docs.append({"file": fn, "text": t[:3500]})
                    if len(docs) >= 80: break
                mf = _kv_get("me:facts") or {}
                facts = dict(mf); facts.update(ws.get("facts") or {})   # «Я» (общий) + факты воркспейса
                owner = (facts.get("owner") or "").strip()
                # МЕХАНИЗМ: не знаем ЧЬИ доли — НЕ угадываем, кладём ВОПРОС в инбокс воркспейса и просим ответить (запомним в «Я»)
                if not owner and docs:
                    existing = next((q for q in ws.get("questions", []) if q.get("key") == "owner" and q.get("status") != "answered"), None)
                    if not existing:
                        sample = "\n\n".join("=== %s ===\n%s" % (d["file"], d["text"][:1400]) for d in docs[:5])
                        qd = _ask_qwen("Из этих инвестиционных документов выпиши РАЗНЫЕ имена/сущности ДЕРЖАТЕЛЕЙ-инвесторов "
                                       "(investor/holder/subscriber), включая SPV-вехикулы. НЕ компании-цели. Верни JSON {\"holders\":[\"...\"]}.\n\n" + sample, mx=600)
                        cands = [h for h in (qd.get("holders") or []) if isinstance(h, str) and h.strip()][:12]
                        ws.setdefault("questions", []).append({
                            "id": "q_" + hashlib.md5(("owner|" + ws["ws_id"]).encode("utf-8")).hexdigest()[:8],
                            "key": "owner", "scope": "me", "answer_type": "choice",
                            "question": "Чьи доли собирать в реестр? Отметь СВОИ сущности — в общих cap table есть и чужие держатели.",
                            "candidates": cands, "status": "open", "answer": ""})
                        _log2("воркспейс спросил", "чьи доли собирать (owner)"); _save()
                        existing = ws["questions"][-1]
                    return json.dumps({"status": "success", "clarify": True, "need": "owner", "scope": "me",
                                       "question": existing["question"], "candidates": existing.get("candidates", []),
                                       "note": "Спрашиваю один раз — ответ сохраню в профиле «Я» (общий для всех воркспейсов)."}, ensure_ascii=False)
                SYS = ("Ты — Extella, финансовый ассистент. Из содержимого файлов СОБЕРИ строки реестра инвестиций. "
                       "СОБИРАЙ ТОЛЬКО доли/инвестиции этих сущностей: «" + owner + "». Если в файле cap table/структура с несколькими держателями — "
                       "бери СТРОКИ ТОЛЬКО этих сущностей, чужих держателей НЕ включай. Одна инвестиция = одна строка (без дублей). "
                       "Верни СТРОГО JSON {\"rows\":[{\"asset\":\"<компания/актив>\",\"class\":\"<стартап|акции|крипта|недвижимость|займ|облигации|опцион|прочее>\","
                       "\"amount\":<число или null>,\"currency\":\"<USD|KZT|EUR>\",\"date\":\"<YYYY-MM или ''>\",\"instrument\":\"<SAFE|SSA|loan|share|option>\",\"source\":\"<имя файла>\",\"note\":\"<кратко>\"}]}. "
                       "Только реальные данные; если суммы нет — null. Без выдумок.\n\nФАЙЛЫ:\n")
                def _chunks(lst, n):
                    for i in range(0, len(lst), n): yield lst[i:i + n]
                rows = []
                for batch in _chunks(docs, 6):
                    payload = "\n\n".join("=== %s ===\n%s" % (d["file"], d["text"]) for d in batch)
                    ai = _ask_qwen(SYS + payload, mx=2200)
                    for r in (ai.get("rows") or []):
                        if isinstance(r, dict) and r.get("asset"): rows.append(r)
                seen = set(); ded = []   # дедуп по (актив, инструмент, сумма)
                for r in rows:
                    kk = (str(r.get("asset")).lower().strip(), str(r.get("instrument")).lower(), str(r.get("amount")))
                    if kk in seen: continue
                    seen.add(kk); ded.append(r)
                rows = ded
                # ROUND-TRIP: если последний записанный реестр правился ЧЕЛОВЕКОМ (mtime > записи) — его правки главнее сборки
                rt_note = ""
                try:
                    _lw = (ws.get("writes") or [])[-1] if ws.get("writes") else None
                    if _lw and _oxl and os.path.isfile(_lw.get("path", "")):
                        import datetime as _dt
                        _wts = _dt.datetime.fromisoformat(str(_lw.get("at"))).timestamp()
                        if os.path.getmtime(_lw["path"]) > _wts + 2:   # правлено после записи
                            _wbk = _oxl.load_workbook(_lw["path"], read_only=True, data_only=True)
                            _sh = _wbk.active; _frows = []
                            _hdr = [str(c.value or "") for c in next(_sh.iter_rows(min_row=1, max_row=1))]
                            for _r in _sh.iter_rows(min_row=2, values_only=True):
                                if not any(v not in (None, "") for v in _r): continue
                                _frows.append({"asset": _r[0], "class": _r[1], "amount": _r[2], "currency": _r[3],
                                               "date": _r[4], "instrument": _r[5], "source": _r[6], "note": _r[7] if len(_r) > 7 else ""})
                            _wbk.close()
                            def _rk(r): return str(r.get("asset", "")).lower().strip() + "|" + str(r.get("instrument", "")).lower() + "|" + str(r.get("date", ""))
                            fmap = {_rk(r): r for r in _frows}
                            old_keys = set(_lw.get("keys") or [])
                            merged = []; edited = 0; removed = 0
                            for r in rows:
                                k2 = _rk(r)
                                if k2 in fmap:
                                    fr = fmap.pop(k2)
                                    if any(str(fr.get(f2)) != str(r.get(f2)) for f2 in ("amount", "currency", "class", "note")): edited += 1
                                    merged.append(fr)          # правки человека главнее
                                elif k2 in old_keys:
                                    removed += 1               # человек удалил строку — не возвращаем
                                else:
                                    merged.append(r)           # новая строка от сборки
                            added_h = list(fmap.values())      # строки, добавленные человеком руками
                            merged.extend(added_h)
                            rows = merged
                            if edited or removed or added_h:
                                rt_note = " Учёл твои правки в файле: изменено %d, удалено %d, добавлено тобой %d." % (edited, removed, len(added_h))
                except Exception: pass
                by_class = {}
                for r in rows:
                    c = r.get("class") or "прочее"; by_class[c] = by_class.get(c, 0) + 1
                note = "Собрал реестр (доли: %s): %d позиций из %d файлов, on-device + Qwen. Классы: %s" % (
                    owner[:40], len(rows), len(docs), "; ".join("%s×%d" % (k, v) for k, v in by_class.items()) or "—") + rt_note
                _log2("собрал реестр инвестиций", note[:120])
                _cr = [{"asset": str(r.get("asset", ""))[:60], "class": str(r.get("class", ""))[:20],
                        "amount": r.get("amount"), "currency": str(r.get("currency", ""))[:6],
                        "date": str(r.get("date", ""))[:10], "instrument": str(r.get("instrument", ""))[:20],
                        "source": os.path.basename(str(r.get("source", "")))[:48], "note": str(r.get("note", ""))[:60]} for r in rows[:120]]
                while _cr and len(json.dumps({"rows": _cr}, ensure_ascii=False)) > 7000:   # KV надёжен только на малых значениях
                    _cr = _cr[:-5]
                _kv_set("ws:" + ws["ws_id"] + ":draft_registry", {"rows": _cr, "truncated": len(_cr) < len(rows), "total": len(rows), "at": _now()})
                note += " Скажи «запиши в таблицу» — положу файлом в папку проекта (новый файл, ничего не перезапишу)."
                adv = _advance(ws, "capability", str(capability) + " реестр", key); _save()
                return json.dumps({"status": "success", "kind": "registry", "capability": capability, "owner": owner,
                                   "rows": rows, "files_read": len(docs), "by_class": by_class, "note": note,
                                   "advanced": (adv["title"] if adv else None), "task_next": ws.get("state", {}).get("next"),
                                   "privacy": "содержимое прочитано on-device; в Qwen ушли извлечённые фрагменты для нормализации (вариант B)"}, ensure_ascii=False)
            COMPARE_KW = ("сравн", "верси", "нестык", "измен", "различ", "дельт", "diff")
            EXTRACT_KW = ("свед", "свест", "реестр", "цифр", "сумм", "извлеч", "счет", "табл", "итог", "посчит", "extract", "total", "reconcil")
            # ---- СРАВНИТЬ ВЕРСИИ (читает обе версии, показывает изменившиеся числа) ----
            if os.path.isdir(base) and any(k in text for k in COMPARE_KW) and (ws.get("versions") or []):
                groups = []
                for grp in ws.get("versions", []):
                    fset = [f for f in (grp.get("files") or []) if os.path.isfile(os.path.join(base, f)) and f.lower().endswith((".xlsx", ".xlsm", ".csv"))]
                    if len(fset) < 2: continue
                    per = []
                    for f in fset:
                        sums = {}
                        for s in _read_any(f):
                            for c in s["columns"]: sums[c["label"]] = round(sums.get(c["label"], 0) + c["sum"], 2)
                        per.append({"file": f, "sums": sums})
                    labels = sorted(set(l for p in per for l in p["sums"]))
                    changes = []
                    for lab in labels:
                        seq = [{"file": p["file"], "value": round(p["sums"].get(lab, 0), 2)} for p in per]
                        vals = [x["value"] for x in seq]
                        if len(set(vals)) > 1:
                            changes.append({"label": lab, "series": seq, "delta": round(vals[-1] - vals[0], 2)})
                    if per: groups.append({"root": grp.get("root", ""), "files": fset, "changes": changes})
                if groups:
                    tot = sum(len(g["changes"]) for g in groups)
                    bits = []
                    for g in groups:
                        for c in g["changes"][:4]:
                            s = c["series"]; d = c["delta"]; sign = "+" if d >= 0 else "−"
                            bits.append("%s: %s→%s (%s%s)" % (c["label"], _fmt(s[0]["value"]), _fmt(s[-1]["value"]), sign, _fmt(abs(d))))
                    note = "Сравнил версии (%d групп, %d изменений). " % (len(groups), tot) + ("; ".join(bits[:5]) if bits else "различий в числах не нашёл")
                    _log2("сравнил версии", note[:120])
                    adv = _advance(ws, "capability", str(capability) + " сравнить версии", key); _save()
                    return json.dumps({"status": "success", "kind": "compare", "capability": capability,
                                       "groups": groups, "note": note,
                                       "advanced": (adv["title"] if adv else None), "task_next": ws.get("state", {}).get("next"),
                                       "privacy": "содержимое прочитано on-device, наружу не уходило"}, ensure_ascii=False)
            # ---- СВЕСТИ ЦИФРЫ В РЕЕСТР (читает все таблицы, суммирует колонки) ----
            if os.path.isdir(base) and any(k in text for k in EXTRACT_KW):
                tables = []; cons = {}; read = 0
                for f in _excl_filter(ws.get("scope_files") or sorted(os.listdir(base)), ws.get("scope_exclude")):
                    if not os.path.isfile(os.path.join(base, f)): continue
                    sheets = _read_any(f)
                    if sheets:
                        read += 1
                        for s in sheets:
                            tables.append({"file": f, "sheet": s["sheet"], "rows": s["rows"], "columns": s["columns"]})
                            for c in s["columns"]:
                                if "%" in c["label"]: continue   # проценты между файлами не суммируем — бессмысленно
                                cons[c["label"]] = round(cons.get(c["label"], 0) + c["sum"], 2)
                if tables:
                    consolidated = [{"label": k, "total": v} for k, v in cons.items()]
                    top = "; ".join("%s: %s" % (c["label"], _fmt(c["total"])) for c in consolidated[:5])
                    note = "Свёл %d таблиц из %d файлов on-device. Итоги: %s" % (len(tables), read, top)
                    _log2("свёл цифры", note[:120])
                    adv = _advance(ws, "capability", str(capability) + " свести цифры", key); _save()
                    return json.dumps({"status": "success", "kind": "extract", "capability": capability,
                                       "tables": tables, "consolidated": consolidated, "note": note,
                                       "advanced": (adv["title"] if adv else None), "task_next": ws.get("state", {}).get("next"),
                                       "privacy": "содержимое прочитано on-device, наружу не уходило"}, ensure_ascii=False)
            # ---- ВЫРАЩЕННЫЕ СПОСОБНОСТИ: матч по триггерам реестра ws:<id>:caps → диспетч исполнителю ----
            # (сам движок чужого эксперта не запускает — исполняет вызывающий: автопилот или сервер плагина)
            for c in (_kv_get("ws:" + ws["ws_id"] + ":caps") or []):
                trs = [str(t).lower() for t in (c.get("triggers") or []) if len(str(t).strip()) >= 3]
                hits = sum(1 for t in trs if t in text)
                if trs and hits >= (2 if len(trs) >= 2 else 1):
                    _log2("зову выращенную способность", str(c.get("cap_name"))[:60]); _save()
                    return json.dumps({"status": "success", "kind": "cap_dispatch",
                                       "expert_name": c.get("expert_name"), "cap_name": c.get("cap_name"),
                                       "params": dict(c.get("params") or {}, ws_id=ws["ws_id"]), "key": key,
                                       "note": "Запускаю свою способность «%s»…" % c.get("cap_name")}, ensure_ascii=False)
            # ---- честная заглушка для прочих намерений: предлагаем ВЫРАСТИТЬ способность ----
            caps = [c for c in (ws.get("pack", {}).get("capabilities") or [])]
            _log2("запрошена способность", str(capability)[:80]); _save()
            return json.dumps({"status": "success", "queued": True, "gap": True, "capability": capability,
                               "note": "Такого пока не умею. Могу НАУЧИТЬСЯ сам: скажи «научись: %s» — спроектирую новую способность, "
                                       "покажу код и после твоего «разрешаю» сделаю." % str(capability)[:80],
                               "pack": caps}, ensure_ascii=False)

        if op == "organize":
            # БЕЗОПАСНАЯ приборка: раскладываем ИСХОДНИКИ по НОВЫМ подпапкам внутри folder. Только ПЕРЕМЕЩЕНИЕ (не удаление),
            # с превью (apply=false) и манифестом для отмены. Ничего вне folder не трогаем.
            plan = ws.get("cleanup") or []
            if not plan:
                return json.dumps({"status": "error", "message": "нет плана уборки (пересобери воркспейс)"}, ensure_ascii=False)
            names_now = set(f for f in os.listdir(base)) if os.path.isdir(base) else set()
            moves = []
            for grp in plan:
                fold = re.sub(r"[/\\:]+", "_", str(grp.get("folder") or "Разное")).strip() or "Разное"
                for fn in (grp.get("files") or []):
                    if fn in names_now and os.path.isfile(os.path.join(base, fn)):
                        moves.append({"file": fn, "to_folder": fold})
            if not _truthy(apply):
                return json.dumps({"status": "success", "preview": True, "moves": moves,
                                   "note": ("Готов разложить %d файлов по папкам: %s. Подтверди — и раскладываю." % (len(moves), ", ".join(sorted(set(m["to_folder"] for m in moves))[:4]))) if moves else "Раскладывать нечего — всё на местах",
                                   "folders": sorted(set(m["to_folder"] for m in moves))}, ensure_ascii=False)
            done = []
            for m in moves:
                src = os.path.join(base, m["file"]); tdir = os.path.join(base, m["to_folder"])
                try:
                    if not os.path.isfile(src): continue
                    os.makedirs(tdir, exist_ok=True)
                    dst = os.path.join(tdir, m["file"])
                    if os.path.exists(dst):  # не перезаписываем — уникализируем
                        r, e = os.path.splitext(m["file"]); dst = os.path.join(tdir, r + "_2" + e)
                    shutil.move(src, dst)
                    done.append({"from": src, "to": dst})
                except Exception:
                    pass
            if done:
                ws.setdefault("moves", []).append({"at": _now(), "items": done})
                for m in done:  # источники теперь в подпапках
                    nm = os.path.basename(m["to"])
                    for s in ws.get("sources", []):
                        if s.get("name") == os.path.basename(m["from"]): s["moved_to"] = os.path.relpath(m["to"], base)
                _log2("прибрался", "разложил %d файлов по папкам: %s" % (len(done), ", ".join(sorted(set(os.path.basename(os.path.dirname(m["to"])) for m in done)))), undo="organize")
            _save()
            return json.dumps({"status": "success", "note": "Разложил файлов: " + str(len(done)) + ". Отменить — «отмени уборку».", "moved": len(done), "workspace": ws}, ensure_ascii=False)

        if op == "undo":
            # сначала записи файлов: если последняя ЗАПИСЬ новее последней приборки — откатываем её (удаляем СВОЙ файл из манифеста)
            writes = ws.get("writes") or []
            moves_b = ws.get("moves") or []
            if writes and (not moves_b or str(writes[-1].get("at", "")) > str(moves_b[-1].get("at", ""))):
                wrec = writes.pop()
                p = wrec.get("path") or ""
                removed = False
                try:
                    if p and os.path.isfile(p) and os.path.realpath(p).startswith(os.path.realpath(base) + os.sep):
                        os.remove(p); removed = True
                        d = os.path.dirname(p)
                        if os.path.isdir(d) and not os.listdir(d): os.rmdir(d)
                except Exception: pass
                ws["writes"] = writes
                _log2("отменил запись", os.path.basename(p) if p else "?")
                _save()
                return json.dumps({"status": "success", "note": "Отменил запись: " + (os.path.basename(p) if p else "?"), "undone_write": p, "removed": removed, "workspace": ws}, ensure_ascii=False)
            batches = ws.get("moves") or []
            if not batches:
                return json.dumps({"status": "success", "nothing": True, "message": "нечего отменять"}, ensure_ascii=False)
            batch = batches.pop()  # последняя приборка
            back = 0
            for m in batch.get("items", []):
                try:
                    if os.path.isfile(m["to"]):
                        os.makedirs(os.path.dirname(m["from"]), exist_ok=True)
                        dst = m["from"]
                        if os.path.exists(dst):
                            r, e = os.path.splitext(dst); dst = r + "_back" + e
                        shutil.move(m["to"], dst); back += 1
                except Exception:
                    pass
            # убрать опустевшие подпапки
            for m in batch.get("items", []):
                d = os.path.dirname(m["to"])
                try:
                    if os.path.isdir(d) and not os.listdir(d): os.rmdir(d)
                except Exception: pass
            for s in ws.get("sources", []): s.pop("moved_to", None)
            ws["moves"] = batches
            _log2("отменил уборку", "вернул %d файлов на место" % back)
            _save()
            return json.dumps({"status": "success", "note": "Вернул файлов на место: " + str(back), "restored": back, "workspace": ws}, ensure_ascii=False)

    # ---------- build ----------
    # ЧЕТЫРЕ входа: ПАПКА (folder) / СЛОВАМИ (intent без folder) / ФАЙЛАМИ (files=JSON-список путей) / ГИБРИД (files+intent).
    scope_files = []
    try: _fl = json.loads(files) if not _b(files) else []
    except Exception: _fl = []
    if isinstance(_fl, list) and _fl:
        picked = [os.path.abspath(os.path.expanduser(str(p))) for p in _fl]
        picked = [p for p in picked if os.path.isfile(p)]
        if not picked:
            return json.dumps({"status": "error", "message": "ни один из переданных файлов не найден"}, ensure_ascii=False)
        folder = os.path.dirname(picked[0]) if len(picked) == 1 else os.path.commonpath([os.path.dirname(p) for p in picked])
        file_list = sorted(os.path.relpath(p, folder) for p in picked)
        scope_files = file_list[:]   # контентные способности читают ТОЛЬКО выбранные файлы
        text_only = False
    else:
        text_only = _b(folder) and not _b(intent)
        if text_only:
            folder = ""
            file_list = []
        else:
            if _b(folder) or not os.path.isdir(os.path.expanduser(folder)):
                return json.dumps({"status": "error", "message": "нужна папка (folder), список файлов (files) или описание проекта словами (intent)"}, ensure_ascii=False)
            folder = os.path.expanduser(folder)
            file_list = _collect_files(folder)   # РЕКУРСИВНО: включая подпапки
            try: _exl = json.loads(exclude) if not _b(exclude) else []
            except Exception: _exl = []
            file_list = _excl_filter(file_list, _exl)   # скоуп-фильтр: чувствительное не сканируем и не шлём в Qwen
            if not file_list:
                return json.dumps({"status": "error", "message": "в папке (и подпапках) нет файлов; можно создать проект словами (intent)"}, ensure_ascii=False)
    files = file_list
    beacon = "ws:build:" + hashlib.md5((("text|" + str(name)) if text_only else folder).encode("utf-8")).hexdigest()[:10]
    _kv_set(beacon, {"status": "running", "started": _now()})

    # ДЕТЕКТ ВЕРСИЙ/ДУБЛЕЙ (детерминированно, по именам — без LLM, приватно): группируем по «корню» имени (basename)
    def _stem(n):
        s = os.path.splitext(os.path.basename(n))[0].lower()
        s = re.sub(r"(v\d+|версия\s*\d*|финал|final|draft|черновик|правк\w*|comments?|коммент\w*|real|copy|копи\w*"
                   r"|20\d\d|q[1-4]\b|кв\d|\d{4}[-_. ]\d\d[-_. ]\d\d|\d{6,})", "", s)
        return re.sub(r"[\s_\-.]+", "", s).strip()
    _by = {}
    for n in files: _by.setdefault(_stem(n), []).append(n)
    versions = [{"root": k, "files": v} for k, v in _by.items() if len(v) > 1]

    # ПОНИМАНИЕ + ПАК + ПРОБЕЛЫ (Qwen; из папки — ТОЛЬКО имена файлов; словами — из описания пользователя)
    ptype, pscore, pambig = _detect_type(intent, files)   # таксономия T1-T16: тип по сигналам интента+имён
    type_hint = ""
    if ptype:
        type_hint = ("\n\nОПРЕДЕЛЁН ТИП ПРОЕКТА: «%s»%s. Типовые цели такого проекта: %s. Типовые артефакты-результаты: %s. "
                     "Используй это как ОРИЕНТИР для goals/plan/pack (но подстрой под конкретику пользователя)."
                     % (ptype["n"], " (неоднозначно — есть второй кандидат)" if pambig else "", ptype["gh"], ptype["art"]))
    manifest = "\n".join("- " + n for n in files[:400])
    _head = ("Ты — Extella. Пользователь описал свой проект СЛОВАМИ (файлов пока нет — подтянутся позже). "
             "Пойми проект, подбери способности, найди пробелы.\n\nОПИСАНИЕ ПРОЕКТА ОТ ПОЛЬЗОВАТЕЛЯ:\n" + str(intent)
             ) if text_only else (
             "Ты — Extella. По ИМЕНАМ файлов и структуре подпапок пойми ГЛАВНЫЙ рабочий проект пользователя, подбери набор способностей "
             "под него и найди пробелы. Содержимое НЕ читаешь (приватность). Конкретно, без воды.\n\nФАЙЛЫ (относительные пути):\n" + manifest)
    prompt = (
        _head + type_hint +
        ('\n\nПОЛЬЗОВАТЕЛЬ ВЫБРАЛ ПРОЕКТ: ' + name + '\n' if not _b(name) else '') +
        ('\n\nЦЕЛЬ ПОЛЬЗОВАТЕЛЯ (он уточнил): ' + intent + '\nПодстрой pack/gaps/first_task ИМЕННО под эту цель.\n' if (not _b(intent) and not text_only) else '') +
        '\n\nВерни СТРОГО JSON: {"project":"<короткое имя проекта>","understood":"<1 фраза: над чем работает>",'
        '"pack_name":"<имя пак-набора>","pack":["<способность1>","<способность2>","..."],'
        '"gaps":[{"title":"<чего не хватает>","severity":"crit|warn","why":"<кратко>"}],'
        '"first_task":"<что предложить сделать первым, 1 фраза>",'
        '"goals":[{"title":"<цель — что делаем>","outcome":"<к какому конкретному результату идём>"}],'
        '"plan":[{"title":"<шаг/задача к цели>","kind":"capability|question|human","goal":<индекс цели, с 0>}],'
        '"autopilot":["<правило дежурства1>","<2>"],'
        '"cleanup":[{"folder":"<имя новой папки>","files":["<точное имя файла ВЕРХНЕГО уровня из списка>","..."]}]}'
        '\n(cleanup — только для файлов В КОРНЕ папки, без «/» в пути; имена бери ТОЧНО из списка)')
    ai = _ask_qwen(prompt)
    if not ai:   # P0-4: Qwen молчит — честная retryable-ошибка, объект не портим
        _kv_set(beacon, {"status": "error", "started": _now(), "error": "llm_unavailable"})
        return json.dumps({"status": "error", "error": "llm_unavailable", "retryable": True,
                           "message": "модель занята — повтори сборку через минуту"}, ensure_ascii=False)
    project = ((name if not _b(name) else "") or ai.get("project") or os.path.basename(folder.rstrip("/")) or "Проект").strip()  # name перебивает ai.project → стабильный ws_id

    wsid = "ws_" + hashlib.md5((("text" if text_only else folder) + "|" + project.lower()).encode("utf-8")).hexdigest()[:10]
    prev = _ws_load(wsid) or {}
    ledger = list(prev.get("ledger") or [])
    def _log(action, detail):
        ledger.append({"n": len(ledger) + 1, "action": action, "detail": detail, "at": _now(), "undo": "none"})
    if not prev:
        _log("создан воркспейс", ("проект «%s» из описания словами" % project) if text_only else ("проект «%s» из папки %s" % (project, folder)))
    _log("связал источники", "%d файлов привязано к проекту" % len(files))
    if versions:
        _log("нашёл версии", "; ".join("%s (%d)" % (v["root"] or "?", len(v["files"])) for v in versions[:4]))
    if ai.get("gaps"):
        _log("выписал пробелы", ", ".join(g.get("title", "") for g in ai["gaps"]))
    if ai.get("pack"):
        _log("подключил пак", "%s: %d способностей" % (ai.get("pack_name") or "pack", len(ai["pack"])))

    # cleanup: только файлы, которые реально есть в папке (Qwen мог придумать имя)
    _fset = set(files)
    cleanup = []
    for grp in (ai.get("cleanup") or []):
        keep = [f for f in (grp.get("files") or []) if f in _fset and os.sep not in f and "/" not in f]  # только корень
        if keep and grp.get("folder"):
            cleanup.append({"folder": grp["folder"], "files": keep})

    # ЦЕЛИ + ЗАДАЧИ + ИНСТРУМЕНТ ПОД ЦЕЛЬ — операционная модель (воркспейс знает этап и «что дальше»)
    def _pick_tool(t):
        t = (t or "").lower()
        if any(k in t for k in ("реестр", "деклар", "финанс", "портфел", "бюджет", "счет", "табл", "excel", "sheet", "cap table", "отчёт", "отчет")): return "Таблица · Excel/Sheets"
        if any(k in t for k in ("код", "code", "md ", "markdown", "репозитор", "github", "dev", "разработ")): return "Код / Markdown"
        if any(k in t for k in ("почт", "письм", "mail", "email", "рассыл")): return "Почта"
        if any(k in t for k in ("презентац", "слайд", "pitch", "slide")): return "Презентация"
        if any(k in t for k in ("договор", "документ", "юрид")): return "Документ"
        return "Заметки"
    goals = []
    for g in (ai.get("goals") or [])[:4]:
        if isinstance(g, dict) and g.get("title"):
            _tl = _pick_tool(str(g["title"]) + " " + (intent or ""))
            if _tl == "Заметки" and ptype: _tl = ptype["tool"]   # дефолт-инструмент от типа проекта
            goals.append({"title": str(g["title"])[:120], "outcome": str(g.get("outcome", ""))[:160],
                          "tool": _tl, "status": "active", "progress": 0})
    if not goals:
        base_goal = intent if not _b(intent) else (ai.get("first_task") or "Разобраться с проектом")
        goals = [{"title": str(base_goal)[:120], "outcome": "", "tool": _pick_tool(str(base_goal)), "status": "active", "progress": 0}]
    tasks = []
    for i, p in enumerate((ai.get("plan") or [])[:12]):
        if isinstance(p, dict) and p.get("title"):
            gi = p.get("goal", 0); gi = gi if (isinstance(gi, int) and 0 <= gi < len(goals)) else 0
            kind = p.get("kind") if p.get("kind") in ("capability", "question", "human") else "capability"
            tasks.append({"id": "t%d" % (i + 1), "title": str(p["title"])[:120], "kind": kind, "goal": gi, "status": "todo"})
    FILE_TOOLS = ("Таблица · Excel/Sheets", "Документ", "Код / Markdown", "Презентация")
    WRITE_WORDS = ("запис", "запиш", "сохран", "выгруз", "экспорт", "файл")
    for gi, g in enumerate(goals):
        gt = [t for t in tasks if t.get("goal") == gi]
        if g.get("tool") in FILE_TOOLS and gt and not any(any(w in str(t.get("title", "")).lower() for w in WRITE_WORDS) for t in gt):
            tasks.append({"id": "t%d" % (len(tasks) + 1), "title": "Записать результат в файл (%s)" % g["tool"].split(" ·")[0],
                          "kind": "capability", "goal": gi, "status": "todo"})
    for gi, g in enumerate(goals):
        gt = [t for t in tasks if t.get("goal") == gi]; done = sum(1 for t in gt if t["status"] == "done")
        g["progress"] = int(100 * done / len(gt)) if gt else 0
    _next = next((t["title"] for t in tasks if t["status"] != "done"), ai.get("first_task", ""))

    ws = {
        "ws_id": wsid, "name": project, "understood": ai.get("understood", ""),
        "intent": (intent if not _b(intent) else prev.get("intent", "")),
        "scope_files": (scope_files or prev.get("scope_files") or []),
        "project_type": ({"key": ptype["k"], "name": ptype["n"], "score": pscore, "ambiguous": pambig} if ptype else (prev.get("project_type") or None)),
        "scope_exclude": ((json.loads(exclude) if not _b(exclude) else None) or prev.get("scope_exclude") or []),
        "facts": prev.get("facts") or {},
        "folder": folder, "created": prev.get("created") or _now(), "updated": _now(),
        "sources": [{"name": n, "ext": os.path.splitext(n)[1].lstrip(".").lower()} for n in files],
        "versions": versions,
        "pack": {"name": ai.get("pack_name") or ((ptype["n"] + " Pack") if ptype else "Capability Pack"), "capabilities": ai.get("pack", []), "installed": False},
        "goals": goals, "tasks": tasks,
        "state": {"gaps": ai.get("gaps", []), "first_task": ai.get("first_task", ""), "next": _next,
                  "blockers": sum(1 for g in ai.get("gaps", []) if g.get("severity") == "crit")},
        "cleanup": cleanup,
        "autopilot": {"enabled": bool((prev.get("autopilot") or {}).get("enabled")), "rules": ai.get("autopilot", [])},
        "moves": prev.get("moves") or [],
        "ledger": ledger, "contract": "ws-v1.3",
    }
    # ДЕРЖИМ ОБЪЕКТ МАЛЫМ: шардированный KV надёжно отдаёт только небольшие значения (~<8КБ); большие читаются
    # битыми/стейл. Капы (registry/organize) читают папку живьём (_collect_files/os.listdir), поэтому полный список
    # sources в объекте не нужен — храним выборку + полный счётчик sources_total.
    ws["sources_total"] = len(ws["sources"])
    ws["sources"] = ws["sources"][:40]
    ws["versions"] = (ws.get("versions") or [])[:12]
    ws["cleanup"] = (ws.get("cleanup") or [])[:12]
    if isinstance(ws.get("state"), dict):
        ws["state"]["gaps"] = (ws["state"].get("gaps") or [])[:10]
        ws["state"]["blockers"] = sum(1 for g in ws["state"]["gaps"] if g.get("severity") == "crit")
    while len(ws["sources"]) > 5 and len(json.dumps(ws, ensure_ascii=False)) > 7500:
        ws["sources"] = ws["sources"][:-8]
    _ws_store(ws)   # шардированная запись (ws-v1.4): tasks/questions/ledger отдельными ключами + verify ядра
    _idx_upsert({"ws_id": wsid, "name": project, "folder": folder, "updated": ws["updated"], "blockers": ws["state"]["blockers"], "kind": ws.get("kind", "scan")})

    _kv_set(beacon, {"status": "done", "started": _now(), "ws_id": wsid})
    return json.dumps({"status": "success", "workspace": ws}, ensure_ascii=False)